# -*- coding: utf-8 -*-
"""Cross-document narrative and statistical-wording consistency check.

    python scripts/check_narrative_consistency.py \
        --manuscript <final>.docx --cover-letter <final>.docx

Verifies that the manuscript, cover letter, REPRODUCIBILITY_AUDIT.md and
README.md tell the same story, and that no statement exceeds what the stored
statistical results support. This is a *wording* audit; numerical agreement
between the documents and the result CSVs is checked by
`scripts/audit_consistency.py`.

Exit code is non-zero if any check fails.
"""
from __future__ import annotations

import argparse
import os
import re
import sys

import pandas as pd

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def docx_text(path: str) -> str:
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for tb in d.tables:
        for row in tb.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def norm(t: str) -> str:
    """Collapse whitespace and strip Markdown block-quote markers, so that a
    phrase split across lines (or quoted with '>') is still found."""
    t = re.sub(r"(?m)^\s*>\s?", "", t)
    return re.sub(r"\s+", " ", t)


# Claims that the stored statistics do not support, anywhere.
BANNED = {
    "equivalence claim": r"\binterchangeable\b|\bequally effective\b|\bidentical performance\b",
    "unsupported superiority": r"statistically superior|significantly outperformed XGBoost|"
                               r"Persistence (was|is) superior(?! )|the most reliable model",
    "causal overreach": r"cannot form a reliable decision boundary|too few [a-z ]*to learn|"
                        r"small enough that a boosted-tree",
    "seed/station confusion": r"(less|more|comparatively) stable\b",
    "operational judgement": r"operationally (in)?sufficient|far from operational",
    "single-command overclaim": r"single (command|script) (regenerates|reproduces) every",
    "checkpoint overclaim": r"trained[- ]model artefacts?",
}
# Sentences that legitimately contain a banned word because they *deny* the claim.
ALLOWED_CONTEXT = [
    "does not demonstrate that the models are equivalent",
    "does not establish equivalence",
    "does not establish that two models are equivalent",
    "we do not claim the architectures are equivalent",
    "not that persistence is superior",
    "not that Persistence was demonstrably superior",
]


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--manuscript", required=True)
    ap.add_argument("--cover-letter", required=True)
    args = ap.parse_args()

    docs = {
        "manuscript": norm(docx_text(args.manuscript)),
        "cover letter": norm(docx_text(args.cover_letter)),
        "REPRODUCIBILITY_AUDIT.md": norm(open(os.path.join(ROOT, "REPRODUCIBILITY_AUDIT.md"),
                                              encoding="utf-8").read()),
        "README.md": norm(open(os.path.join(ROOT, "README.md"), encoding="utf-8").read()),
    }
    MS, CL, AUD, RM = (docs["manuscript"], docs["cover letter"],
                       docs["REPRODUCIBILITY_AUDIT.md"], docs["README.md"])
    ALL = " ".join(docs.values())

    results = []

    def ck(name, ok, detail=""):
        results.append((name, bool(ok), detail))

    # ---- banned wording -------------------------------------------------
    violations = []
    for label, pat in BANNED.items():
        for dname, text in docs.items():
            for m in re.finditer(pat, text, re.I):
                seg = text[max(0, m.start() - 120): m.end() + 120]
                if any(a.lower() in seg.lower() for a in ALLOWED_CONTEXT):
                    continue
                violations.append(f"[{dname}] {label}: …{seg}…")
    ck("no unsupported ranking / equivalence / causal / capability claims",
       not violations, " || ".join(violations))

    # ---- statistical stance ---------------------------------------------
    ck("Persistence vs XGBoost reported as non-significant after Holm",
       "p = 0.068" in MS and "p = 0.068" in CL and "0.068" in AUD)
    ck("XGBoost framed as 'did not demonstrate an improvement'",
       "did not demonstrate an improvement over" in MS
       and "did not improve on persistence" in CL)
    ck("onset XGBoost vs LSTM reported as non-significant",
       "p = 0.301" in MS and "p = 0.301" in CL)
    ck("non-significance explicitly distinguished from equivalence",
       "does not demonstrate that the models are equivalent" in MS
       and "does not establish equivalence" in CL)
    ck("Holm correction family described as task x metric",
       "Within each prediction task and evaluation metric" in MS
       and "and evaluation metric" in AUD and "and evaluation metric" in RM)
    ck("station SD described as between-station variability, not stability",
       "between-station variability" in MS
       and "not a measure of run-to-run stability" in MS)
    ck("low-prevalence explanation is associative, not causal",
       "may partly reflect" in MS and "consistent with the scarcity" in MS)

    # ---- core narrative --------------------------------------------------
    core = "does not necessarily translate into useful fog-onset early-warning skill"
    ck("core message present in manuscript (abstract + conclusions)",
       MS.count(core) >= 2)
    ck("core message present in cover letter", core in CL)
    ck("core message present in audit", core in AUD)
    ck("overall and onset framed as separate tasks",
       "as separate tasks" in MS and "as separate tasks" in CL)
    ck("contribution framed as evaluation framework, not a new best model",
       "evaluation framework and a diagnostic finding" in MS
       and "rather than a new best-performing model" in MS)
    ck("Persistence's structural inability to detect onset stated",
       "unable to detect 0→1 fog onset" in MS and "cannot issue an onset warning" in MS)
    ck("deep-learning result scoped, not generalised",
       "not be read as evidence that deep learning is unsuitable" in MS
       and ("in the overall T+1 task evaluated here" in MS
            or "under this data configuration" in MS))

    # ---- repository capability -------------------------------------------
    ck("trained models declared not redistributed",
       "trained model files are not redistributed" in MS.lower()
       and "trained model files are not redistributed" in CL.lower())
    ck("document tools declared outside the reproduction pipeline",
       "not part of `scripts/reproduce_all.py`" in AUD)
    ck("manuscript-level audit declared a separate invocation",
       "separate manual invocation" in AUD and "--manuscript" in RM)

    # ---- numbers untouched -----------------------------------------------
    met = pd.read_csv(os.path.join(ROOT, "results", "metrics", "station_metrics.csv"))
    f1o = met[met.task == "overall"].pivot_table(index="station", columns="model", values="f1")
    ck("headline numbers still match results/metrics/station_metrics.csv",
       f"{f1o['Persistence'].mean():.3f} ± {f1o['Persistence'].std(ddof=1):.3f}" == "0.578 ± 0.122"
       and "0.578 ± 0.122" in MS and "0.507 ± 0.210" in MS)

    for name, ok, detail in results:
        print(f"{'PASS' if ok else 'FAIL'}  {name}")
        if not ok and detail:
            print(f"        {detail[:600]}")
    n_fail = sum(1 for _, ok, _ in results if not ok)
    print(f"\nPASS = {len(results) - n_fail}    FAIL = {n_fail}")
    return 1 if n_fail else 0


if __name__ == "__main__":
    raise SystemExit(main())
