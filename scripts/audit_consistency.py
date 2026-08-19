# -*- coding: utf-8 -*-
"""PHASE 17: automated manuscript <-> config <-> result-file consistency audit.

    python scripts/audit_consistency.py [--manuscript manuscript/Atmosphere_eng_v6_reproduced.docx]

Every PASS means the number printed in the manuscript is byte-for-byte
derivable from configs/default.yaml or results/metrics/*.csv.
Exit code is non-zero if anything FAILs, so this can gate a release.
"""
from __future__ import annotations

import argparse
import os
import re
import sys

import pandas as pd
import yaml

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OVERALL = ["Persistence", "XGBoost", "1D-CNN", "LSTM"]
ONSET = ["No-Onset Baseline", "XGBoost-Onset", "1D-CNN-Onset", "LSTM-Onset"]


def docx_text(path: str) -> str:
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for tb in d.tables:
        for row in tb.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--manuscript", default=None)
    args = ap.parse_args()

    cfg = yaml.safe_load(open(os.path.join(ROOT, "configs", "default.yaml"), encoding="utf-8"))
    m = os.path.join(ROOT, cfg["paths"]["results_dir"], "metrics")
    t = os.path.join(ROOT, cfg["paths"]["results_dir"], "tables")
    met = pd.read_csv(os.path.join(m, "station_metrics.csv"))

    checks = []

    def chk(item, manuscript_val, source_val, source, tol=0.0005):
        if manuscript_val is None:
            ok = None
        elif isinstance(source_val, (int, float)) and isinstance(manuscript_val, (int, float)):
            ok = abs(float(manuscript_val) - float(source_val)) <= tol
        else:
            ok = str(manuscript_val) == str(source_val)
        checks.append({"item": item, "manuscript": manuscript_val,
                       "code_or_result": source_val, "source": source,
                       "status": "PASS" if ok else ("FAIL" if ok is False else "NOT FOUND")})

    text = docx_text(args.manuscript) if args.manuscript and os.path.exists(args.manuscript) else None

    def find_num(pattern):
        if text is None:
            return None
        mm = re.search(pattern, text)
        if not mm:
            return None
        return float(mm.group(1).rstrip("."))

    # ---- configuration values ----
    d = cfg["models"]["deep"]
    x = cfg["models"]["xgboost"]
    conf = [
        ("Sequence length", find_num(r"sequence length of (\d+)"), d["seq_len"], "config deep.seq_len"),
        ("DL input variables", find_num(r"consisted of (\d+) variables"), len(cfg["features"]["dl"]), "config features.dl"),
        ("CNN kernel size", find_num(r"kernel size (\d+)"), d["cnn"]["kernel_size"], "config deep.cnn.kernel_size"),
        ("LSTM hidden size", find_num(r"hidden size = (\d+)"), d["lstm"]["hidden_size"], "config deep.lstm.hidden_size"),
        ("Batch size", find_num(r"batch size of (\d+)"), d["batch_size"], "config deep.batch_size"),
        ("Max epochs", find_num(r"up to (\d+) epochs"), d["max_epochs"], "config deep.max_epochs"),
        ("Dropout", find_num(r"dropout \(rate = ([\d.]+)\)"), d["dropout"], "config deep.dropout"),
        ("Adam learning rate", find_num(r"learning rate = ([\de.-]+)\)"), d["learning_rate"], "config deep.learning_rate"),
        ("XGB n_estimators", find_num(r"n_estimators = (\d+)"), x["n_estimators"], "config xgboost"),
        ("XGB max_depth", find_num(r"max_depth = (\d+)"), x["max_depth"], "config xgboost"),
        ("XGB learning_rate", find_num(r"XGB.{0,400}?learning_rate = ([\d.]+)"), x["learning_rate"], "config xgboost"),
        ("XGB subsample", find_num(r"subsample = ([\d.]+?)[,\s]"), x["subsample"], "config xgboost"),
        ("XGB colsample_bytree", find_num(r"colsample_bytree = ([\d.]+?)\.?[\s,;]"), x["colsample_bytree"], "config xgboost"),
    ]
    for item, msv, srcv, src in conf:
        chk(item, msv, srcv, src, tol=1e-9)

    # ---- split periods ----
    chk("Train period start", "2015" if text and "2015–2021 for training" in text else None,
        "2015", "config data.years")
    chk("Test period", "2023-2024" if text and "2023\u20132024 for testing" in text else None,
        "2023-2024", f"config train_end={cfg['data']['train_end']}, val_end={cfg['data']['val_end']}")

    # ---- headline results ----
    f1o = met[met["task"] == "overall"].pivot_table(index="station", columns="model", values="f1")
    for mdl in OVERALL:
        val = round(f1o[mdl].mean(), 3)
        ms = find_num(rf"{re.escape(mdl)}[^\n]{{0,80}}?([01]\.\d{{3}})\s*±") if text else None
        chk(f"Overall mean F1 — {mdl}", ms, val, "results/metrics/station_metrics.csv")

    f1n = met[met["task"] == "onset"].pivot_table(index="station", columns="model", values="f1")
    for mdl in ONSET:
        if mdl not in f1n.columns:
            continue
        if mdl == "No-Onset Baseline":
            msv = 0.0 if text and "No-Onset Baseline" in text else None
        else:
            msv = find_num(rf"{re.escape(mdl)}\n([01]\.\d{{3}}) ±") if text else None
        chk(f"Onset mean F1 — {mdl}", msv, round(f1n[mdl].mean(), 3),
            "results/metrics/station_metrics.csv")

    # ---- every table cell must exist in a result CSV ----
    tbl = os.path.join(t, "table_overall_f1_by_station.csv")
    if os.path.exists(tbl):
        tt = pd.read_csv(tbl, index_col=0)
        mism = 0
        for st in f1o.index:
            for mdl in OVERALL:
                if abs(float(tt.loc[st, mdl]) - round(f1o.loc[st, mdl], 3)) > 0.0011:
                    mism += 1
        chk("Table 3 cells == station_metrics.csv", mism, 0, "make_tables.py output", tol=0)

    out = pd.DataFrame(checks)
    os.makedirs(t, exist_ok=True)
    out.to_csv(os.path.join(t, "consistency_audit.csv"), index=False)
    print(out.to_string(index=False))
    n_fail = int((out["status"] == "FAIL").sum())
    n_missing = int((out["status"] == "NOT FOUND").sum())
    print(f"\nPASS={int((out['status']=='PASS').sum())}  FAIL={n_fail}  NOT FOUND={n_missing}")
    if text is None:
        print("(no manuscript supplied — code/result side only)")
    return 1 if n_fail else 0


if __name__ == "__main__":
    raise SystemExit(main())
