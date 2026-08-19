# -*- coding: utf-8 -*-
"""Rewrite the cover letter to match the reproduced results.

    python update_cover_letter.py

Reads the same result CSVs as update_manuscript.py; the original file is left
untouched and a new version is written alongside it.
"""
from __future__ import annotations

import os

import docx
import pandas as pd

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(HERE)   # repository root
M = os.path.join(REPO, "results", "metrics")
T = os.path.join(REPO, "results", "tables")

SRC = os.environ.get("COVER_LETTER_TEMPLATE",
                     os.path.join(HERE, "cover_letter_template.docx"))
OUT = os.path.join(HERE, "cover_letter_v5_persistence_lit.docx")

met = pd.read_csv(os.path.join(M, "station_metrics.csv"))
hh = pd.read_csv(os.path.join(T, "table_persistence_vs_xgboost.csv"))
shap_g = pd.read_csv(os.path.join(M, "shap_global.csv"))
sig = pd.read_csv(os.path.join(T, "table_significance_tests.csv"))


def sg(task, a, b, metric="f1"):
    m = sig[(sig["task"] == task) & (sig["metric"] == metric)]
    hit = m[(m["model_a"] == a) & (m["model_b"] == b)]
    if len(hit):
        return hit.iloc[0]
    return m[(m["model_a"] == b) & (m["model_b"] == a)].iloc[0]


def pstr(p):
    return "p < 0.001" if p < 0.001 else f"p = {p:.3f}"

f1o = met[met["task"] == "overall"].pivot_table(index="station", columns="model", values="f1")
f1n = met[met["task"] == "onset"].pivot_table(index="station", columns="model", values="f1")
prn = met[met["task"] == "onset"].pivot_table(index="station", columns="model", values="pr_auc")
ONSET = ["No-Onset Baseline", "XGBoost-Onset", "1D-CNN-Onset", "LSTM-Onset"]

pers_wins = int(hh["persistence_beats_xgb_tuned"].sum())
pers_wins_def = int((hh["persistence_f1"] > hh["xgboost_f1_default_0.5"]).sum())
xgb_best_onset = int((f1n[ONSET].idxmax(axis=1) == "XGBoost-Onset").sum())


def ms(s):
    return f"{s.mean():.3f} ± {s.std(ddof=1):.3f}"


REPLACEMENTS = {
    3: ("Motivation and contribution. Persistence is an established reference point in "
        "very-short-term low-visibility forecasting: Cornejo-Bueno et al. (Symmetry 2020, 12, "
        "1045) analysed low-visibility persistence at Valladolid Airport and evaluated hourly "
        "prediction with a naïve persistence operator alongside machine-learning methods. "
        "Short-term fog studies nevertheless do not always benchmark learning-based models "
        "against an explicit persistence baseline, and they do not always evaluate fog onset "
        "separately from routine fog persistence. Our study is deliberately framed as a "
        "methodological and diagnostic contribution rather than a claim of state-of-the-art "
        "performance, and we do not claim novelty for the use of a persistence baseline "
        "itself. Using ten years (2015–2024) of hourly ASOS observations from twelve stations "
        "across the central-northern Korean Peninsula, we make three points that we believe "
        "are useful to both researchers and forecasters:"),

    4: ("1. We separate overall T+1 fog prediction from 0→1 onset prediction and show that "
        "strong overall one-hour-ahead fog-state forecasting performance does not necessarily "
        "translate into useful fog-onset early-warning skill — the baseline that is hardest to "
        "beat on the overall task is precisely the one that cannot issue an onset warning at "
        "all. Our central argument is therefore that the two must be specified and evaluated "
        "as separate tasks."),

    2: ("We are pleased to submit our manuscript entitled “Short-term Fog Forecasting and "
        "Onset Detection Using Multi-station ASOS Observations in Korea: A Comparison of "
        "Persistence, XGBoost, and Deep Learning Models” for consideration for publication "
        "in Atmosphere."),

    5: (f"2. Consistent with earlier low-visibility work, we find persistence to be a very "
        f"strong baseline for very short-range fog prediction, and we verify this across "
        f"twelve Korean ASOS stations. It attained the highest mean F1-score of "
        f"the four models tested ({ms(f1o['Persistence'])} versus "
        f"{ms(f1o['XGBoost'])} for XGBoost), outperforming XGBoost at {pers_wins} of the "
        f"twelve stations after validation-based threshold tuning and at all "
        f"{pers_wins_def} stations at the default decision threshold, with the margin "
        f"widening where fog is rare — a pattern consistent with the scarcity of positive "
        f"training cases at those stations. Across twelve paired station observations the "
        f"Persistence–XGBoost difference is not significant after Holm correction "
        f"({pstr(sg('overall','Persistence','XGBoost')['p_holm'])}), so we claim only that "
        f"XGBoost did not improve on persistence, not that persistence is superior. Both, "
        f"however, exceeded the two deep models significantly "
        f"({pstr(sg('overall','XGBoost','LSTM')['p_holm'])} or better) under this "
        f"experimental configuration. Persistence "
        f"is nevertheless structurally incapable of detecting onset. This clarifies what "
        f"learning-based models do and do not add, and we consider a clearly negative "
        f"result of this kind worth reporting."),

    6: (f"3. We show that onset detection remains a genuinely hard, largely unsolved task in "
        f"absolute terms for all models tested. A tree-based model "
        f"(XGBoost, mean F1 {f1n['XGBoost-Onset'].mean():.3f}) achieved the highest mean "
        f"F1-score and was best at {xgb_best_onset} of "
        f"the twelve stations, ahead of LSTM ({f1n['LSTM-Onset'].mean():.3f}) and 1D-CNN "
        f"({f1n['1D-CNN-Onset'].mean():.3f}); on mean PR-AUC the LSTM was marginally ahead "
        f"({prn['LSTM-Onset'].mean():.3f} vs. {prn['XGBoost-Onset'].mean():.3f}); neither "
        f"difference is statistically significant "
        f"({pstr(sg('onset','XGBoost-Onset','LSTM-Onset')['p_holm'])} on F1, "
        f"{pstr(sg('onset','XGBoost-Onset','LSTM-Onset','pr_auc')['p_holm'])} on PR-AUC), so "
        f"we report the two as comparable within the statistical resolution of our data "
        f"rather than ranked; we do not claim the architectures are equivalent. SHAP analysis "
        f"indicates "
        f"that predictions are driven by physically interpretable and regionally consistent "
        f"variables (current visibility, dew point depression, prior visibility and relative "
        f"humidity)."),

    7: ("Honest scope. We state the limitations explicitly in the manuscript. The comparison "
        "between XGBoost and the deep-learning models reflects practical model "
        "configurations rather than a controlled architecture comparison, although we did "
        "equalise the temporal splits, the evaluation cases and the information horizon "
        "available to every model. Model differences were tested with two-sided Wilcoxon "
        "signed-rank tests over the twelve stations with Holm correction, and we have "
        "deliberately withheld ranking claims that those tests do not support. With only "
        "twelve paired stations, statistical power is limited, a non-significant result does "
        "not establish equivalence, and the stations may not constitute fully independent "
        "spatial units since they lie within the same regional meteorological system. The "
        "reported between-station standard deviations reflect regional heterogeneity rather "
        "than estimation uncertainty or run-to-run stability. Very-low-prevalence stations "
        "carry substantial "
        "uncertainty, and the deep-learning results depend on random initialisation; we "
        "report a fixed-seed run and quantify the seed-to-seed variability in the "
        "repository. We believe that reporting these negative and cautionary findings "
        "honestly is itself a contribution, as it helps prevent over-interpretation of "
        "aggregate scores in operational fog nowcasting."),

    8: ("This manuscript is original, has not been published elsewhere, and is not under "
        "consideration by any other journal. All authors have approved the submission and "
        "declare no conflicts of interest. The ASOS observation data used are publicly "
        "available from the Korea Meteorological Administration. To support verification, we "
        "release a public repository containing the complete pipeline, configuration files "
        "and result files. A documented pipeline regenerates all reported metrics, "
        "statistical test results, tables and figures from the raw observations, and the "
        "deep-learning models are checked structurally before each run to confirm that the "
        "1D-CNN and LSTM are what they are described to be. Trained model files are not "
        "redistributed, as the pipeline reproduces them from the raw data."),
}


def main() -> int:
    doc = docx.Document(SRC)
    idx = [i for i, p in enumerate(doc.paragraphs) if p.text.strip()]
    for n, new_text in REPLACEMENTS.items():
        p = doc.paragraphs[idx[n]]
        for r in p.runs[1:]:
            r.text = ""
        if p.runs:
            p.runs[0].text = new_text
        else:
            p.add_run(new_text)
    doc.save(OUT)
    print(f"Wrote {OUT}  ({len(REPLACEMENTS)} paragraphs replaced)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
