# -*- coding: utf-8 -*-
"""검토용 국문본 생성 (manuscript + cover letter).

    python make_korean_review.py

영문 정본 Atmosphere_eng_v8_final.docx / cover_letter_v4_final.docx 의 번역본이며,
모든 수치는 영문본과 동일하게 fog-forecast-repro/results/ 의 CSV에서 직접 읽는다.
번역본에서 새로운 주장·해석·수치를 만들지 않는다.
"""
from __future__ import annotations

import os
import zipfile

import docx
import pandas as pd
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(HERE)   # repository root
M = os.path.join(REPO, "results", "metrics")
T = os.path.join(REPO, "results", "tables")
FIGDIR = os.environ.get("FIGDIR", os.path.join(REPO, "results", "figures_docx"))

KO_FONT = "맑은 고딕"

# ----------------------------------------------------------------- results
met = pd.read_csv(os.path.join(M, "station_metrics.csv"))
prev = pd.read_csv(os.path.join(M, "fog_prevalence.csv")).set_index("station")
trans = pd.read_csv(os.path.join(T, "table_transition_counts_test.csv"))
shap_g = pd.read_csv(os.path.join(M, "shap_global.csv"))
errv = pd.read_csv(os.path.join(M, "error_visibility_rows.csv"))
hh = pd.read_csv(os.path.join(T, "table_persistence_vs_xgboost.csv"))
split = pd.read_csv(os.path.join(M, "split_summary.csv"))
sig = pd.read_csv(os.path.join(T, "table_significance_tests.csv"))

STATIONS_KO = {
    "Baengnyeongdo": "백령도", "Daegwallyeong": "대관령", "Paju": "파주",
    "Incheon": "인천", "Ganghwa": "강화", "Dongducheon": "동두천",
    "Chuncheon": "춘천", "Taebaek": "태백", "Cheorwon": "철원",
    "Sokcho": "속초", "Seoul": "서울", "Inje": "인제",
}
ORDER = list(STATIONS_KO)
OVERALL = ["Persistence", "XGBoost", "1D-CNN", "LSTM"]
ONSET = ["No-Onset Baseline", "XGBoost-Onset", "1D-CNN-Onset", "LSTM-Onset"]

f1o = met[met.task == "overall"].pivot_table(index="station", columns="model",
                                             values="f1").reindex(ORDER)
f1n = met[met.task == "onset"].pivot_table(index="station", columns="model",
                                           values="f1").reindex(ORDER)
prn = met[met.task == "onset"].pivot_table(index="station", columns="model",
                                           values="pr_auc").reindex(ORDER)


def ms(sr):
    return f"{sr.mean():.3f} ± {sr.std(ddof=1):.3f}"


def sg(task, a, b, metric="f1"):
    m = sig[(sig.task == task) & (sig.metric == metric)]
    h = m[(m.model_a == a) & (m.model_b == b)]
    return (h if len(h) else m[(m.model_a == b) & (m.model_b == a)]).iloc[0]


def pk(p):
    return "p < 0.001" if p < 0.001 else f"p = {p:.3f}"


pers_wins = int(hh["persistence_beats_xgb_tuned"].sum())
xgb_wins = 12 - pers_wins
xgb_def_mean = hh["xgboost_f1_default_0.5"].mean()
xgb_best_onset = int((f1n[ONSET].idxmax(axis=1) == "XGBoost-Onset").sum())
best_onset = f1n[ONSET].idxmax(axis=1)
losers_ko = [STATIONS_KO[s] for s in ORDER if best_onset[s] != "XGBoost-Onset"]

prev = prev.assign(pct=prev["fog_prevalence_test"] * 100)
pv = prev.sort_values("pct", ascending=False)
mean_prev = prev["pct"].mean()
ratio = pv["pct"].iloc[0] / pv["pct"].iloc[-1]

spw = met[met.model == "XGBoost"]["scale_pos_weight"].mean()
spw_on = met[met.model == "XGBoost-Onset"]["scale_pos_weight"].mean()
thr_o = met[met.model == "XGBoost"]["threshold"]
thr_n = met[met.model == "XGBoost-Onset"]["threshold"]

tr = {r["transition"]: r for _, r in trans.iterrows()}
onset_pct = float(tr["0→1"]["proportion_percent"])

tp = errv[errv.error_type == "TP"]["vis_tplus1_10m"].median() / 100
fn = errv[errv.error_type == "FN"]["vis_tplus1_10m"].median() / 100
fp = errv[errv.error_type == "FP"]["vis_tplus1_10m"].median() / 100
n_tp = int((errv.error_type == "TP").sum())
n_fn = int((errv.error_type == "FN").sum())
n_fp = int((errv.error_type == "FP").sum())

cov = split.pivot(index="station", columns="split", values="n_rows")[["train", "val", "test"]]
cov["total"] = cov.sum(axis=1)
cov["pct"] = cov["total"] / (10 * 365.25 * 24) * 100
covs = cov.sort_values("pct")


# ----------------------------------------------------------------- helpers
def set_ko(run):
    run.font.name = KO_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KO_FONT)


def para(doc, text, size=10.5, bold=False, align=None, space_after=6,
         italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    set_ko(r)
    return p


def heading(doc, text, level=1):
    sizes = {0: 16, 1: 13, 2: 11.5}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level < 2 else 10)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.size = Pt(sizes[level])
    r.bold = True
    set_ko(r)
    return p


def table(doc, rows, widths=None, header=True, size=9):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 or i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(str(val))
            r.font.size = Pt(size)
            r.bold = (i == 0 and header)
            set_ko(r)
            if widths:
                cell.width = Cm(widths[j])
    return t


def caption(doc, text):
    para(doc, text, size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


def figure(doc, img, cap, width=15.0):
    path = os.path.join(FIGDIR, img)
    if not os.path.exists(path):
        para(doc, f"[그림 누락: {img}]", size=9, italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(path, width=Cm(width))
    caption(doc, cap)


def new_doc():
    d = docx.Document()
    s = d.sections[0]
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    for attr, v in (("top_margin", 2.2), ("bottom_margin", 2.2),
                    ("left_margin", 2.4), ("right_margin", 2.4)):
        setattr(s, attr, Cm(v))
    st = d.styles["Normal"]
    st.font.name = KO_FONT
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), KO_FONT)
    return d


def notice(doc, lines):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.cell(0, 0)
    c.text = ""
    for i, line in enumerate(lines):
        p = c.paragraphs[0] if i == 0 else c.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        set_ko(r)
    doc.add_paragraph()


# ================================================================ MANUSCRIPT
def build_manuscript(out_path):
    d = new_doc()

    para(d, "[검토용 국문본]", size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    heading(d, "다관측소 ASOS 관측자료를 이용한 한반도 단기 안개 예측 및 발생 탐지: "
               "지속성(Persistence)·XGBoost·딥러닝 모델 비교", level=0)
    para(d, "우석문 1, 김인영 1,*", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(d, "1 국방대학교 사이버·컴퓨터과학과, 논산 33021, 대한민국 / "
            "* 교신저자: inyoungkim@korea.ac.kr",
         size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    notice(d, [
        "■ 본 문서는 지도교수 검토를 위한 국문 번역본입니다. 학술지 투고 정본은 "
        "Atmosphere_eng_v8_final.docx 이며, 표현이 다를 경우 영문본이 우선합니다.",
        "■ 본 번역본의 모든 수치는 영문본과 동일하게 fog-forecast-repro/results/metrics/ 및 "
        "results/tables/ 의 CSV에서 자동으로 읽어 생성되었습니다. 번역 과정에서 수치를 "
        "직접 입력하지 않았습니다.",
        "■ 참고문헌 목록은 원문(영문) 표기를 그대로 유지했습니다.",
    ])

    # ---------------- 초록 ----------------
    heading(d, "초록", level=1)
    para(d,
         "안개는 시정을 급격히 저하시켜 항공·도로·해상 운항과 군사 활동에 직접적인 영향을 "
         "미치며, 급격한 발생과 소멸 특성 때문에 단기 예측이 여전히 어렵다. 본 연구는 "
         "한반도 중북부 12개 ASOS 관측소의 10년(2015–2024) 시간별 자료를 이용하여 단기 안개 "
         "예측을 평가하되, 전체 T+1(1시간 후) 안개 예측과 0→1 안개 발생(onset) 탐지를 서로 "
         "분리하여 평가하였다. 입력변수로는 기상변수, lag·rolling 통계량, 이슬점차(dew point "
         "depression), 지면–대기 온도차를 사용하였다. 네 모델은 모두 동일한 시험 사례 집합에서 "
         "평가되었다. "
         f"전체 T+1 예측에서는 학습이 필요 없는 지속성 기준모델이 가장 높은 평균 F1을 "
         f"기록하였고({ms(f1o['Persistence'])}), XGBoost({ms(f1o['XGBoost'])}), "
         f"LSTM({ms(f1o['LSTM'])}), 1D-CNN({ms(f1o['1D-CNN'])})이 그 뒤를 이었다. 지속성은 "
         f"검증셋 기반 임계값 조정 후 12개 중 {pers_wins}개 관측소에서, 기본 임계값에서는 "
         f"12개 전 관측소에서 XGBoost보다 우수하였다. 다만 지속성–XGBoost 차이는 다중비교 "
         f"보정 후 유의하지 않았으며({pk(sg('overall','Persistence','XGBoost')['p_holm'])}), "
         f"두 모델 모두 두 딥러닝 모델보다는 유의하게 높았다"
         f"({pk(sg('overall','XGBoost','LSTM')['p_holm'])} 이하). 그러나 지속성은 구조적으로 "
         "안개 발생을 탐지할 수 없다. "
         f"발생 탐지에서는 모든 학습 모델의 절대 성능이 낮았으며, XGBoost가 가장 높은 평균 "
         f"F1({ms(f1n['XGBoost-Onset'])})을 기록하고 12개 중 {xgb_best_onset}개 관측소에서 "
         f"최고였으나, LSTM({ms(f1n['LSTM-Onset'])}) 및 "
         f"1D-CNN({ms(f1n['1D-CNN-Onset'])})과 비교할 때 XGBoost와 LSTM은 F1"
         f"({pk(sg('onset','XGBoost-Onset','LSTM-Onset')['p_holm'])})과 "
         f"PR-AUC({prn['LSTM-Onset'].mean():.3f} 대 {prn['XGBoost-Onset'].mean():.3f}) 모두에서 "
         "통계적으로 구분되지 않았다. SHAP 분석 결과 현재 시정, 이슬점차, 상대습도가 관측소 "
         "전반에 걸쳐 가장 영향력이 큰 변수로 나타났다. "
         "이러한 결과는 전체 안개 상태 예측 성능이 우수하다고 해서 그것이 곧 유용한 안개 발생 "
         "조기경보 능력으로 이어지지는 않는다는 점을 보여준다. 전체 과제에서 가장 이기기 어려웠던 "
         "기준모델은 애초에 발생 경보를 낼 수 없으며, 발생 경보가 가능한 학습 모델들은 절대 "
         "성능이 제한적이다. 따라서 전체 안개 상태 예측과 안개 발생 탐지는 서로 다른 과제로 "
         "규정하고 별도로 평가해야 한다.")
    para(d, "주제어: 안개 예측; 안개 발생; 지속성 기준모델; XGBoost; SHAP; 시정; 지역 이질성; 조기경보",
         size=9.5, bold=True, space_after=10)

    # ---------------- 1. 서론 ----------------
    heading(d, "1. 서론", level=1)
    para(d, "안개는 지표 부근 수증기가 미세 물방울로 응결되어 수평 시정을 저하시키는 대기 현상으로, "
            "일반적으로 시정 1 km 미만으로 정의된다. 안개로 인한 저시정은 도로·항공·해상 교통뿐 "
            "아니라 기동, 감시·정찰, 비행 운영과 관련된 군사적 의사결정에도 직접 영향을 미친다. "
            "안개는 짧은 시간 안에 형성·소멸할 수 있으므로 단기 안개 예측은 단순한 기상정보 제공을 "
            "넘어 조기경보와 작전적 의사결정을 지원하는 중요한 문제이다 [4,13,17].")
    para(d, "그러나 안개 예측은 여전히 어려운 문제이다. 안개 형성은 기온, 이슬점온도, 상대습도, "
            "풍속, 지면 냉각, 대기 안정도, 수분 공급, 경계층 과정 등 다수 기상 인자들의 비선형 "
            "상호작용에 의해 결정되며, 지형·고도·해안으로부터의 거리·해륙 열대비 같은 국지 "
            "지리조건에도 민감하다. 삼면이 바다이고 산악 지형이 발달한 한반도는 안개 발생의 "
            "시공간 변동성이 크다. Lee 등[13]은 한반도 안개가 혼합안개, 해무, 복사안개, 전선안개, "
            "증기안개, 활승안개 등 다양한 유형으로 나타나며 지역·계절에 따라 특성이 다르다고 "
            "보고하였다. Kim 등[20]도 남서부 지역의 산악·내륙·해안 간 안개 형성 기구가 다르며 "
            "이러한 차이가 기계학습 모델의 예측 성능에 영향을 준다고 보고하였다. 따라서 안개 예측 "
            "모델은 단일 관측소의 평균 성능만으로 평가할 것이 아니라, 서로 다른 안개 특성을 갖는 "
            "여러 지역에서 검증되어야 한다.")
    para(d, "전통적으로는 수치예보(NWP) 모델과 물리 기반 시정 모수화가 안개·저시정 예측에 "
            "사용되어 왔다. 이 접근은 복사냉각, 난류 혼합, 미세물리 과정, 경계층 발달 등 안개 "
            "형성과 관련된 물리 과정을 표현할 수 있다는 장점이 있다. 그러나 안개는 국지적이고 "
            "짧은 시간 규모에서 발생하는 경우가 많아, 현업 NWP 모델의 시공간 해상도로는 발생 시점과 "
            "위치를 정확히 재현하기 어렵다. 특히 시정은 액체수함량, 물방울 수농도, 입경분포, "
            "에어로졸, 소산계수 같은 미세물리 변수의 영향을 받는데, 이들은 관측과 모수화 모두에서 "
            "불확실성이 크다 [10,15,17]. Kamangir 등[22]도 현업 NWP 모델이 미세물리 및 행성경계층 "
            "과정을 직접 해석하지 않고 모수화에 의존하기 때문에 안개의 발생·위치·시점 예측에 "
            "구조적 어려움이 있다고 지적하였다.")
    para(d, "최근 안개·저시정 예측 연구에서는 Random Forest, Gradient Boosting, XGBoost, LightGBM, "
            "CatBoost 같은 트리 기반 앙상블과 LSTM, CNN, ResNet-LSTM 같은 딥러닝 모델이 폭넓게 "
            "적용되어 왔다 [1,3–5,9,14,20–22]. 트리 기반 앙상블은 정형 기상자료에서 변수 간 "
            "비선형 관계와 상호작용을 효과적으로 학습할 수 있고 상대적으로 적은 자료에서도 "
            "경쟁력 있는 성능을 낼 수 있다. Schütz 등[1]은 지상·위성·결합 자료를 이용한 "
            "XGBoost 기반 시정 예측을 수행하면서, 전체 시정 예측만이 아니라 안개 형성과 소멸 "
            "사례를 분리하여 평가할 필요성을 강조하였다. Penov와 Guerova[3]는 공항 시정 추정에서 "
            "Random Forest와 LSTM을 비교하여 이슬점차와 안개 안정도 지수가 중요한 예측변수임을 "
            "보고하였다. 국내에서는 Kim 등[14]이 세종·부산 스마트시티 관측자료를 이용해 Random "
            "Forest 및 DNN 기반 시정 추정을 수행하여, 내륙과 해안 도시 간 안개 특성과 모델 성능이 "
            "다름을 보였다.")
    para(d, "딥러닝 기반 연구는 시간적·공간적 구조를 직접 학습할 수 있다는 점에서 안개 예측에 "
            "적용되어 왔다. Sun 등[5]은 ResNet-LSTM 기반 해무 예측 모델을 제안하고 focal loss로 "
            "클래스 불균형을 완화하였다. Kamangir 등[22]은 NWP 산출물과 위성 기반 해수면온도를 "
            "결합한 다중척도 3D-CNN 모델 FogNet을 제안하여, 6·12·24시간 선행 시정 등급 예측에서 "
            "현업 HREF 앙상블을 능가하였다. 그러나 이러한 접근은 NWP 산출물·위성·해양자료 등 "
            "대규모 고차원 격자 입력을 요구하여, 자료 구성과 계산 비용 측면에서 일반적인 지상관측 "
            "기반 현업 환경과는 차이가 있다. 반면 해석 가능하고 계산 효율이 높은 트리 기반 모델은 "
            "ASOS 같은 지상관측 기반 지역 단기 예측에서 여전히 실용적인 대안이다.")
    para(d, "클래스 불균형 또한 안개 예측의 핵심 난제이다. 안개는 대부분의 시간대에 발생하지 않는 "
            "희소 현상이므로 전체 시계열에서 안개 사례 비율이 매우 낮다. 특히 비안개 상태에서 "
            "다음 시각에 안개로 전이되는 0→1 발생 사례는 전체 안개 사례보다도 더 희소하다. 이러한 "
            "불균형에서는 모델이 높은 정확도를 달성하면서도 실제 안개 발생을 제대로 탐지하지 못할 "
            "수 있다. Shin 등[19]은 국내 주요 도시 시정 초단기예측에서 클래스 불균형과 시간적 분포 "
            "변화가 성능 저하의 주요 원인이라고 지적하였다. Kim 등[20] 역시 국내 ASOS 기반 안개 "
            "빈도 예측에서 학습자료 중 안개 사례가 5% 미만이며 이러한 심한 불균형이 트리 기반 "
            "분류기 학습을 저해한다고 보고하였다. 따라서 안개 예측 모델은 단순 정확도가 아니라 "
            "정밀도(precision), 재현율(recall), F1-score 등 탐지 성능을 직접 반영하는 지표로 "
            "평가되어야 한다.")
    para(d, "초단기 안개 예측에서는 현재 시정 상태와 안개 지속성 또한 중요한 인자이다. 안개는 일단 "
            "형성되면 일정 시간 지속되는 경향이 있으므로, 현재 또는 직전 시정은 다음 시각 시정의 "
            "강력한 예측변수가 될 수 있다. Peláez-Rodríguez 등[21]은 스페인 산악도로 저시정 "
            "예측에서 과거 시정값, 즉 지속성 관련 변수가 핵심 예측변수임을 보였다. 그러나 지속성은 "
            "조기경보 관점에서 구조적 한계를 갖는다. 현재 시각이 비안개이면 지속성 모델은 다음 "
            "시각도 비안개로 예측하므로, 비안개에서 안개로의 전이인 0→1 안개 발생을 탐지할 수 "
            "없다. 다시 말해 지속성이 전체 T+1 안개 예측에서 좋은 성능을 보이더라도, 이는 운영상 "
            "더 중요한 안개 발생 예측 성능이 높다는 것을 의미하지 않는다.")
    para(d, "최근 일부 연구는 안개를 정적 분류 문제가 아니라 상태 전이 문제로 다루기 시작하였다. "
            "Schütz 등[1]은 안개 형성과 소멸을 별도 사건으로 간주하여 XGBoost 기반 시정 예측을 "
            "평가하고, 전체 성능 지표가 안개 전이 사건에 대한 예측 성능을 가릴 수 있음을 보였다. "
            "Han 등[18]은 인천항과 해운대의 해무 소멸을 별도 예측 문제로 다루며 소멸 시점 예측의 "
            "운영적 중요성을 강조하였다. 그러나 한반도 다수 지상관측소를 대상으로 전체 T+1 안개 "
            "예측과 0→1 안개 발생 예측을 명시적으로 분리하고, 지속성 기준모델의 구조적 한계와 "
            "학습 모델의 발생 탐지 성능을 함께 분석한 연구는 아직 제한적이다.")
    para(d, "국내 국지 안개 예측 연구도 수행되어 왔다. Jeong과 Kim[2]은 강원도 태백의 2015–2024년 "
            "기상자료를 이용하여 포화수증기압·상대습도·이슬점 관계를 물리 손실함수에 반영한 "
            "물리정보신경망(PINN) 기반 국지 안개 예측 모델을 제안하였고, 안개 탐지에서 ANN, "
            "XGBoost, LightGBM 및 군 경험식보다 우수한 성능을 보였다. 다만 이 연구는 단일 관측소를 "
            "대상으로 하였고 전체 다음시각 안개 예측과 0→1 발생 예측을 분리 평가하지 않았다. 국내 "
            "다른 연구들도 해무, 스마트시티 시정, 재분석 기반 안개 빈도, 도시 시정 초단기예측 등을 "
            "다루었으나, 다수 ASOS 관측소를 대상으로 조기경보 관점에서 발생 예측을 분리 평가한 "
            "연구는 여전히 부족하다 [12,14,19,20].")
    para(d, "기계학습 기반 안개 예측 모델이 실제 의사결정에 활용되려면 예측 성능만큼 해석 가능성도 "
            "중요하다. 예보관이나 운영자가 모델의 판단 근거를 이해하지 못하면 결과를 신뢰하고 "
            "활용하기 어렵다. 이에 따라 SHAP 같은 설명가능 AI 기법이 기상·기후 분야에서 블랙박스 "
            "모델의 예측 근거를 해석하는 데 활용되어 왔다. Yang 등[6]은 기상·기후 예측에서 해석 "
            "가능성이 모델 신뢰 구축, 오류 진단, 물리적 이해 확보에 중요하다고 설명하였다.")
    para(d, "이러한 한계를 바탕으로 본 연구는 한반도 중북부 12개 ASOS 관측소를 대상으로 단기 안개 "
            "예측을 수행한다. 안개는 대상 시각의 시정 1 km 이하로 정의하고, 예측 대상은 1시간 후 "
            "안개 발생 여부이다. 본 연구는 모든 상태 전이를 포함하는 전체 T+1 안개 예측과, 비안개 "
            "상태에서 다음 시각 안개로 전이되는 0→1 안개 발생 예측을 명시적으로 구분한다. 모델로는 "
            "지속성, XGBoost, 1D-CNN, LSTM을 비교하며, 분류 임계값은 검증셋에서 F1 기준으로 "
            "선택하고 최종 성능은 독립 시험셋에서 평가한다. 아울러 XGBoost 모델에 SHAP 분석을 "
            "적용하여 주요 예측변수의 기여도와 지역별 변수 중요도 구조를 해석한다.")
    para(d, "본 연구의 주요 기여는 다음과 같다. 첫째, 전체 T+1 안개 예측과 0→1 안개 발생 예측을 "
            "분리함으로써 전체 성능 평가만으로는 조기경보 성능을 충분히 특성화할 수 없음을 보인다. "
            "둘째, 초단기 안개 예측에서 지속성을 강력한 기준모델로 평가하는 동시에 0→1 발생을 "
            "탐지할 수 없는 구조적 한계를 실증한다. 셋째, 동일한 시간 분할, 동일한 평가 사례, 모든 "
            "모델에 동일한 정보 지평을 부여하는 입력 윈도우라는 엄격히 대등한 프로토콜 아래 12개 "
            "ASOS 관측소에서 XGBoost, 1D-CNN, LSTM을 비교하여 지역별 안개 예측 가능성과 관측소 간 "
            "성능 변동성을 분석한다. 넷째, SHAP 분석을 통해 XGBoost 예측이 현재·직전 시정, "
            "이슬점차, 상대습도 등 물리적으로 해석 가능한 변수에 기반하는지 검증한다. 따라서 본 "
            "연구가 의도하는 기여는 새로운 최고성능 모델의 제안이 아니라, 단기 안개 예측 능력을 "
            "어떻게 측정해야 하는가에 관한 평가 프레임워크와 진단적 발견이다. 보고된 지표·표·그림을 "
            "재현하는 데 필요한 모든 코드, 설정파일, 결과파일을 공개한다.")

    # ---------------- 2. 재료 및 방법 ----------------
    heading(d, "2. 재료 및 방법", level=1)
    heading(d, "2.1. 연구 지역 및 자료", level=2)
    para(d, "본 연구는 한반도 중북부 12개 ASOS 관측소를 사용하였다: 백령도, 대관령, 파주, 인천, "
            "동두천, 강화, 춘천, 철원, 태백, 속초, 서울, 인제. 이들 관측소는 도서·해안·내륙·산악 "
            "지형을 함께 포괄하므로 지역별 안개 특성을 비교하기에 적합하다. 백령도는 서해 도서 "
            "관측소로 해양 영향이 클 수 있고, 인천과 강화는 해안 및 수도권 서부 특성을 반영한다. "
            "파주·동두천·철원·춘천은 내륙 및 접경지역 특성을, 대관령과 태백은 고지대 산악 지형에서 "
            "활승안개와 복사냉각의 영향을 반영할 것으로 예상된다. 속초와 인제는 각각 동해안과 인접 "
            "산악 지형의 영향을 받을 수 있다.")
    figure(d, "image1.png", "그림 1. 연구 지역 및 12개 ASOS 관측소 위치. 삽입 지도는 한반도 내 "
                            "연구 지역의 위치를 나타낸다.", 12.5)
    para(d, "기상청(KMA)이 제공하는 시간별 ASOS 자료를 사용하였다. 분석 기간은 2015–2024년이며, "
            "관측소별 자료는 기온, 이슬점온도, 상대습도, 풍속, 현지기압, 지면온도, 시정을 포함한다. "
            "자료는 기상자료개방포털을 통해 취득하였고, 모델 학습 및 평가에 앞서 관측소별·시간별로 "
            "결측과 이상치를 전처리하였다. "
            f"자료 커버리지는 12개 관측소에서 균일하지 않다. 정제 후 사용 가능한 시간별 자료 수는 "
            f"{int(covs['total'].min()):,}건에서 {int(covs['total'].max()):,}건까지 분포하며, 이는 "
            f"2015–2024년 명목 시간자료의 {covs['pct'].min():.0f}–{covs['pct'].max():.0f}%에 "
            f"해당한다. 커버리지가 가장 낮은 곳은 {STATIONS_KO[covs.index[0]]}"
            f"({covs['pct'].iloc[0]:.0f}%), {STATIONS_KO[covs.index[1]]}"
            f"({covs['pct'].iloc[1]:.0f}%), {STATIONS_KO[covs.index[2]]}"
            f"({covs['pct'].iloc[2]:.0f}%)이다. {STATIONS_KO[covs.index[1]]}과 "
            f"{STATIONS_KO[covs.index[2]]}은 결손이 주로 시험 기간에 발생하여 시험셋 사례 수가 다른 "
            f"관측소의 약 절반 수준이며, {STATIONS_KO[covs.index[0]]}은 결손이 주로 학습 기간에 "
            "발생한다. 분할별 관측소 자료 수는 공개 저장소에 수록하였다. 따라서 관측소별 결과는 "
            "이러한 이질성을 고려하여 비교되어야 한다.")
    para(d, "대상 시각의 안개는 수평 시정 1 km 이하로 정의하였다 [13,17]. ASOS 시정 변수는 10 m "
            "단위로 보고되므로 시정 값이 100 이하인 자료를 안개(1), 그 외를 비안개(0)로 표기하였다. "
            "이는 표준 기상학적 안개 정의와 일치한다.")
    para(d, "자료는 시간 순서를 보존하여 분할하였다. 학습 2015–2021년, 검증 2022년, 시험 "
            "2023–2024년이다. 이러한 시간 기반 분할은 시계열 자료에서 미래 정보가 학습 과정으로 "
            "누설되는 것을 방지한다. 검증셋은 모델별 최적 분류 임계값 선택에 사용하였고, 최종 "
            "성능은 독립 시험셋에서 평가하였다.")

    heading(d, "2.2. 전처리 및 파생변수 생성", level=2)
    para(d, "모든 모델은 동일한 원시 관측자료와 동일한 시간 분할을 사용하였다. 입력 표현은 모델 "
            "구조에 따라 달랐다. XGBoost는 28개 변수의 정형 특성집합을 사용하였고, 1D-CNN과 LSTM은 "
            "선택된 기상·시차·물리 파생·시간 변수 14개로 구성된 6시간 시퀀스를 사용하였다. 두 표현 "
            "모두 시각 t에서 끝나므로 시퀀스 윈도우는 t−5부터 t까지를 포함하며, 두 모델군은 동일한 "
            "1시간 예측 지평을 갖는다.")
    para(d, "첫째, 시차(lag) 변수를 생성하였다. 직전 1–3시간의 기온, 습도, 시정 값을 입력변수로 "
            "포함하여 단기 지속성과 추세를 포착하였다. 특히 직전 시정은 안개의 지속과 소멸에 밀접히 "
            "관련되어 T+1 안개 예측의 중요한 단서가 된다. 시차 및 이동 연산은 규칙적인 시간 축 "
            "위에서만 의미가 있으므로, 물리적 타당성 필터를 통과하지 못한 관측값은 삭제하지 않고 "
            "결측 처리한 뒤 각 관측소 시계열을 완전한 시간 격자로 재색인하고 나서 시차·이동 변수를 "
            "계산하였다. 이후 여전히 결측이 남은 행을 제거하였다.")
    para(d, "둘째, 기온과 상대습도에 대해 이동 통계량을 생성하였다. 2시간 이동평균과 이동표준편차를 "
            "{t−1, t} 윈도우에서 계산하였다. 따라서 이 통계량은 예보 발표 시점 이전에 이용 가능한 "
            "관측값만 사용하며 대상 시각 t+1의 정보를 포함하지 않는다.")
    para(d, "셋째, 물리 파생변수를 구성하였다. 이슬점차는 기온과 이슬점온도의 차로 계산하였으며 "
            "값이 작을수록 포화에 가까운 상태를 의미한다. 지면–대기 온도차와 그 1시간 변화량도 "
            "산출하였다. 이들 변수는 안개 형성과 관련된 대기 포화 및 지면 냉각 조건을 표현한다.")
    para(d, "넷째, 시간 및 계절 변수를 구성하였다. 시각, 요일, 주말 여부를 포함하고 월은 "
            "sine/cosine 변환으로 부호화하여 계절 주기성을 연속형으로 표현하였다.")
    para(d, "예측 대상은 시각 t+1의 안개 상태이며, 모든 예측변수는 시각 t 이전에 관측 가능한 "
            "정보로 제한하였다. 시각 t의 현재 시정도 그러한 예측변수 중 하나로, 예보 발표 시점에 "
            "이용 가능하며 모든 학습 모델의 입력에 포함된다. 이는 학습 모델과 지속성 기준모델이 "
            "동일한 정보에 접근하도록 하기 위한 것이다. t+1의 정보를 포함하는 관측값이나 파생변수는 "
            "입력에 포함하지 않았으며, 이 시간 정렬은 특성 목록에 대한 프로그램적 단언(assertion)으로 "
            "강제된다.")

    heading(d, "2.3. 전체 T+1 예측과 발생 예측의 정의", level=2)
    para(d, "본 연구는 두 가지 예측 과제를 정의한다. 첫째는 전체 T+1 안개 예측으로, 모든 시각 t에서 "
            "시각 t의 기상자료를 이용해 t+1의 안개 발생 여부를 예측하며 네 가지 상태 전이를 모두 "
            "포함한다. 둘째는 발생(onset) 예측으로, 현재 시각이 비안개인 사례만을 후보로 제한하고 "
            "다음 시각에 안개가 새로 발생하는지를 예측한다. 즉 현재 안개가 이미 존재하는 사례(1→0, "
            "1→1 전이)는 발생 예측 평가에서 제외된다.")
    caption(d, "표 1. 전체 T+1 예측 및 발생 예측의 전이 유형 정의. 0은 비안개, 1은 안개를 의미한다.")
    table(d, [["현재 상태", "다음 상태", "전이 유형", "전체 T+1 예측", "발생 예측"],
              ["0", "0", "비안개 지속", "포함", "비발생"],
              ["0", "1", "안개 발생", "포함", "발생"],
              ["1", "0", "안개 소멸", "포함", "제외"],
              ["1", "1", "안개 지속", "포함", "제외"]],
          widths=[2.6, 2.6, 3.6, 3.4, 3.0])
    d.add_paragraph()
    para(d, "발생 예측은 전체 T+1 예측보다 어려운 과제이다. 전체 T+1 예측에서는 기존 안개가 다음 "
            "시각까지 지속되는 사례(1→1)가 성능 향상에 기여할 수 있다. 반면 발생 예측은 현재 비안개 "
            "상태에서 다음 시각에 안개가 새로 형성되는 사례(0→1)만을 탐지해야 한다. 따라서 발생 "
            "예측은 운영상 조기경보 관점에서 전체 T+1 예측보다 더 직접적인 관련성을 갖는다.")

    heading(d, "2.4. 비교 모델", level=2)
    para(d, "본 연구는 지속성, XGBoost, 1D-CNN, LSTM 네 모델을 비교하였다. 지속성은 학습이 필요 없는 "
            "기준모델로, 현재 안개 상태를 다음 시각 예측으로 그대로 이월한다. 발생 예측에서는 평가 "
            "후보가 현재 비안개인 사례로 완전히 제한되므로 지속성 모델은 모든 사례를 비발생으로 "
            "예측하게 된다. 본 연구는 이를 No-Onset 기준모델이라 부른다.")
    para(d, "XGBoost는 Chen과 Guestrin[7]이 제안한 그래디언트 부스팅 기반 트리 앙상블 모델이다. "
            "XGBClassifier 구현을 사용하였으며 초매개변수는 n_estimators = 250, max_depth = 5, "
            "learning_rate = 0.03, subsample = 0.6, colsample_bytree = 0.6으로 설정하였다. "
            "scale_pos_weight는 학습자료의 음성/양성 표본 비로 계산하되 관측소와 과제별로 각각 "
            f"산출하였으며, 평균값은 전체 T+1 예측에서 약 {spw:.1f}, 발생 예측에서 약 {spw_on:.1f}"
            "이었다. 전체 T+1 예측과 발생 예측용 XGBoost 모델은 독립적으로 학습하였다.")
    para(d, "1D-CNN과 LSTM은 PyTorch(2.13)로 구현하였으며 입력 시퀀스 길이는 t−5부터 t까지 6시간이다. "
            "시퀀스 입력은 14개 변수로 구성된다: 기본 기상변수 7개(기온, 이슬점온도, 습도, 풍속, "
            "현지기압, 지면온도, 시정), 이슬점차, 지면–대기 온도차, 직전 시각의 시정과 습도, 시각, "
            "sine/cosine 부호화 월. 모든 변수는 학습 분할에서만 적합한 StandardScaler로 표준화하였다. "
            "시퀀스는 시간 분할 이후에 구성하였고, 6개 시각이 정확히 1시간 간격일 때만 윈도우를 "
            "유지하여 어떤 윈도우도 분할 경계나 자료 공백을 가로지르지 않도록 하였다.")
    para(d, "1D-CNN은 시간 축을 따라 적용되는 1차원 합성곱 2개(32→64 채널, 커널 크기 3, same 패딩, "
            "ReLU), 시간 축 전역 평균 풀링, 드롭아웃(0.3), 완전연결층(64→32, ReLU), 출력층(32→1)으로 "
            "구성되며 학습 가능 매개변수는 9,697개이다. LSTM은 단층(은닉 크기 64)의 최종 은닉 상태를 "
            "드롭아웃(0.3), 완전연결층(64→32, ReLU), 출력층(32→1)에 전달하는 구조이며 학습 가능 "
            "매개변수는 22,593개이다. 두 모델 모두 손실함수로 BCEWithLogitsLoss를 사용하고 pos_weight를 "
            "학습자료의 음성/양성 비로 설정하여 클래스 불균형에 대응하였다. 최적화는 Adam(학습률 "
            "1e-3), 배치 크기 512, 최대 15 에폭이며, 검증 F1이 3 에폭 연속 개선되지 않으면 조기 "
            "종료하고 검증 F1이 가장 높았던 에폭의 가중치를 최종 모델로 사용하였다.")
    para(d, "모든 실험 전에 자동 구조 점검이 수행된다. 1D-CNN이 시간 축에서 동작하는 실제 nn.Conv1d "
            "층을 포함하는지, LSTM이 최종 은닉 상태를 분류기에 전달하는 실제 nn.LSTM 모듈을 "
            "포함하는지, 두 모듈 모두에 기울기가 도달하는지, 입력 시퀀스를 역순으로 바꾸면 출력이 "
            "변하는지를 확인하며, 하나라도 실패하면 실험이 실행되지 않는다.")
    para(d, "모든 학습 모델은 동일한 학습/검증/시험 분할을 사용하였다. 시퀀스 모델은 완전한 6시간 "
            "이력이 필요하므로 자료 공백 직후의 소수 시험 행을 점수화할 수 없다. 따라서 지속성과 "
            "XGBoost를 포함한 네 모델 모두 모든 모델이 점수화할 수 있는 동일한 시험 사례 부분집합에서 "
            "평가하여, 보고되는 모든 F1이 동일한 분모와 동일한 양성 클래스 비율을 갖도록 하였다.")

    heading(d, "2.5. 임계값 최적화 및 성능 평가", level=2)
    para(d, "확률을 출력하는 모델(XGBoost, 1D-CNN, LSTM)에 대해 최적 분류 임계값은 검증셋에서 양성 "
            "클래스 F1을 기준으로 선택하였다. 0.05–0.95 구간을 181개 등간격 값으로 탐색하였으며, "
            f"12개 관측소에서 전체 T+1 XGBoost 모델의 선택 임계값은 {thr_o.min():.3f}–"
            f"{thr_o.max():.3f}(중앙값 {thr_o.median():.3f}), 발생 XGBoost 모델은 "
            f"{thr_n.min():.3f}–{thr_n.max():.3f}(중앙값 {thr_n.median():.3f})의 범위였다. 심한 "
            "클래스 불균형과 클래스 가중치 때문에 모델 출력 확률은 보정된 확률이라기보다 운영적 "
            "결정 점수로 기능한다는 점에 유의해야 한다. 즉 임계값 0.94는 안개 발생 확률 94%를 "
            "의미하지 않으며 검증 F1을 최대화하도록 선택된 운영 임계값으로 해석해야 한다. 검증셋에서 "
            "선택된 임계값은 고정한 뒤 시험셋에 적용하였다. 시험셋은 임계값 선택, 모델 선택, 조기 "
            "종료를 비롯한 어떠한 결정에도 사용되지 않았다.")
    para(d, "모델 간 차이는 12개 관측소에 대한 양측 Wilcoxon 부호순위 검정으로 평가하였다. 각 "
            "관측소가 하나의 대응 관측치를 제공하며, 12개의 대응 차이만으로는 정규성 가정을 "
            "정당화하기 어렵고 관측소별 F1 값의 이질성이 크기 때문이다. 각 예측 과제와 평가 지표별로 "
            "모든 쌍별 비교의 p값을 Holm 단계적 절차로 보정하였다. 따라서 보정 계열은 전체 과제 "
            "F1(6개 비교), 발생 과제 F1(3개 비교), 발생 과제 PR-AUC(3개 비교)의 세 가지이다. "
            "효과크기는 대응쌍 rank-biserial 상관 r로 보고하며, 첫 번째 모델이 모든 관측소에서 이길 "
            "때 +1이 된다.")
    para(d, "성능 평가는 양성 클래스의 정밀도, 재현율, F1을 중심으로 하였다. 본 자료에서 안개와 발생 "
            "사례는 매우 불균형하므로, 전체 정확도가 아니라 양성 클래스 탐지 성능을 직접 반영하는 "
            "정밀도·재현율·F1을 주요 평가지표로 사용하였다.")

    heading(d, "2.6. SHAP 기반 모델 해석", level=2)
    para(d, "XGBoost 예측의 근거를 해석하기 위해 SHAP 분석을 수행하였다. SHAP은 Lundberg와 "
            "Lee[8]가 제안한, 게임이론의 Shapley 값 개념에 기반한 설명가능 AI 기법으로 각 입력변수가 "
            "개별 예측에 기여하는 정도를 정량적으로 추정한다. 본 연구는 변수별 평균 절대 SHAP 값을 "
            "계산하여 전역 변수 중요도를 평가하고, 관측소별 평균 |SHAP| 값을 비교하여 지역별 주요 "
            "변수 기여도의 차이를 분석하였다. SHAP 분석은 전체 T+1 XGBoost 모델에 대해 수행하였고 "
            "TP·FP·FN 오차 분석과 함께 해석하였다.")

    # ---------------- 3. 결과 ----------------
    heading(d, "3. 결과", level=1)
    heading(d, "3.1. 지역별 안개 발생률", level=2)
    figure(d, "image2.png", f"그림 2. 관측소별 시험셋 안개 발생률. 12개 관측소 평균 안개 발생률은 "
                            f"{mean_prev:.2f}%이다.")
    para(d, f"시험셋에서 12개 관측소의 평균 안개 발생률은 약 {mean_prev:.2f}%로, 분석 기간 대부분이 "
            f"비안개 조건이며 안개 클래스가 매우 불균형함을 보여준다. 그러나 발생률은 지역별로 크게 "
            f"달랐다. {STATIONS_KO[pv.index[0]]}가 {pv['pct'].iloc[0]:.2f}%로 가장 높았고 "
            f"{STATIONS_KO[pv.index[1]]}({pv['pct'].iloc[1]:.2f}%), "
            f"{STATIONS_KO[pv.index[2]]}({pv['pct'].iloc[2]:.2f}%)가 뒤를 이어 모두 전체 평균을 "
            f"상회하였다. 반면 {STATIONS_KO[pv.index[-1]]}({pv['pct'].iloc[-1]:.2f}%), "
            f"{STATIONS_KO[pv.index[-2]]}({pv['pct'].iloc[-2]:.2f}%), "
            f"{STATIONS_KO[pv.index[-3]]}({pv['pct'].iloc[-3]:.2f}%)는 발생률이 매우 낮았다. "
            f"발생률이 가장 높은 {STATIONS_KO[pv.index[0]]}와 가장 낮은 "
            f"{STATIONS_KO[pv.index[-1]]}를 비교하면 약 {ratio:.0f}배의 차이가 나타난다. 이는 본 "
            "연구의 안개 예측이 단순한 양성/음성 클래스 불균형뿐 아니라 지역 간 발생 빈도의 상당한 "
            "공간적 불균형을 포함함을 의미한다.")

    heading(d, "3.2. 전체 T+1 전이 구조 분석", level=2)
    para(d, f"12개 관측소 시험셋 전체에서 가장 흔한 전이 유형은 비안개 지속(0→0)으로 "
            f"{tr['0→0']['count']:,}건({tr['0→0']['proportion_percent']:.2f}%)이었다. 안개 발생"
            f"(0→1)은 {tr['0→1']['count']:,}건({tr['0→1']['proportion_percent']:.2f}%), 안개 소멸"
            f"(1→0)은 {tr['1→0']['count']:,}건({tr['1→0']['proportion_percent']:.2f}%), 안개 지속"
            f"(1→1)은 {tr['1→1']['count']:,}건({tr['1→1']['proportion_percent']:.2f}%)이었다.")
    rows = [["전이 유형", "설명", "총 건수", "비율"]]
    desc = {"0→0": "비안개 지속", "0→1": "안개 발생", "1→0": "안개 소멸", "1→1": "안개 지속"}
    for _, r in trans.iterrows():
        rows.append([r["transition"], desc[r["transition"]], f"{r['count']:,}",
                     f"{r['proportion_percent']:.2f}%"])
    caption(d, "표 2. 시험셋 전이 유형 요약.")
    table(d, rows, widths=[3.0, 5.0, 3.6, 3.6])
    d.add_paragraph()
    para(d, "이 결과는 전체 T+1 안개 예측 성능을 해석하는 데 중요한 함의를 갖는다. 전체 T+1 "
            "예측에서 지속성 모델은 현재 안개 상태를 다음 시각으로 그대로 이월하므로 1→1 안개 지속 "
            "사례는 쉽게 탐지하지만, 0→1 발생 사례는 구조적으로 탐지할 수 없다.")

    heading(d, "3.3. 전체 T+1 모델 성능의 지역별 비교", level=2)
    figure(d, "image3.png", "그림 3. 시험셋 안개 발생률 순으로 정렬한 관측소별 전체 T+1 안개 예측 "
                            "성능 비교. 지속성, XGBoost, 1D-CNN, LSTM의 안개 클래스 F1을 동일한 "
                            "시험 사례 집합에서 비교하였다.")
    para(d, f"XGBoost, 1D-CNN, LSTM은 검증셋에서 F1을 최대화하는 임계값을 선택한 뒤 고정하여 "
            f"시험셋에서 평가하였다. 이 절차에서 지속성은 12개 중 {pers_wins}개 관측소에서 XGBoost보다 "
            f"우수하였고, 나머지 {xgb_wins}개에서는 XGBoost가 우수하였다. 지속성의 평균 F1"
            f"({f1o['Persistence'].mean():.3f})은 XGBoost({f1o['XGBoost'].mean():.3f})보다 "
            f"{f1o['Persistence'].mean()-f1o['XGBoost'].mean():.3f} 높았다. 안개 발생률이 높은 "
            f"관측소에서는 두 모델이 근접하였으나"
            f"(XGBoost가 {', '.join(STATIONS_KO[s] for s in hh[~hh['persistence_beats_xgb_tuned']]['station'])}에서 근소 우위), "
            f"발생률이 낮은 관측소에서는 격차가 크게 벌어졌다"
            f"(서울 {f1o.loc['Seoul','XGBoost']:.3f} 대 {f1o.loc['Seoul','Persistence']:.3f}; "
            f"인제 {f1o.loc['Inje','XGBoost']:.3f} 대 {f1o.loc['Inje','Persistence']:.3f}). 이는 "
            "해당 관측소에서 양성 학습 사례가 희소하여 견고한 결정경계 학습이 제약된 상황과 부합하는 "
            "양상이다.")
    para(d, f"XGBoost의 관측소 간 표준편차({f1o['XGBoost'].std(ddof=1):.3f})는 지속성"
            f"({f1o['Persistence'].std(ddof=1):.3f})보다 상당히 컸다. 즉 XGBoost는 관측소 간 변동성이 "
            "더 컸다. 이 값은 12개 관측소에 걸친 성능 이질성을 기술하는 것이며, 서로 다른 무작위 "
            "초기화에서의 실행 간 안정성을 측정한 값이 아니다. 12개 관측소에 대한 양측 Wilcoxon "
            f"부호순위 검정에서 이 차이는 Holm 보정 후 유의하지 않았다"
            f"(보정 전 {pk(sg('overall','Persistence','XGBoost')['p_raw'])}, 보정 후 "
            f"{pk(sg('overall','Persistence','XGBoost')['p_holm'])}; rank-biserial "
            f"r = {sg('overall','Persistence','XGBoost')['rank_biserial_r']:.2f}). 따라서 적절한 "
            "해석은 XGBoost가 지속성보다 개선되지 않았다는 것이며, 지속성이 통계적으로 우월함이 "
            "입증되었다는 것은 아니다.")
    rows = [["관측소", "지속성", "XGBoost", "1D-CNN", "LSTM"]]
    for s in ORDER:
        rows.append([STATIONS_KO[s]] + [f"{f1o.loc[s, m]:.3f}" for m in OVERALL])
    rows.append(["평균±표준편차"] + [f"{f1o[m].mean():.3f}±{f1o[m].std(ddof=1):.3f}" for m in OVERALL])
    caption(d, "표 3. 전체 T+1 모델 성능의 지역별 비교(F1, 시험셋 2023–2024).")
    table(d, rows, widths=[3.4, 3.2, 3.2, 3.2, 3.2])
    d.add_paragraph()
    para(d, f"1D-CNN은 관측소 평균 성능이 가장 낮았고({ms(f1o['1D-CNN'])}), LSTM은 1D-CNN보다 "
            f"우수하였으나 XGBoost와 지속성보다는 낮았다({ms(f1o['LSTM'])}). 두 딥러닝 모델 모두 "
            f"어느 관측소에서도 최고 F1을 기록하지 못하였다. 분류 임계값을 기본값 0.5로 고정하면 "
            f"지속성이 12개 전 관측소에서 XGBoost보다 우수하며 XGBoost 평균 F1은 "
            f"{xgb_def_mean:.3f}까지 하락한다. 이는 학습 모델이 지속성 대비 경쟁력을 갖는 것처럼 "
            "보이는 정도가 임계값 최적화 절차에 크게 의존함을 보여준다.")
    para(d, "지속성–XGBoost 비교와 달리, 두 모델과 두 딥러닝 모델을 가르는 격차는 Holm 보정 후에도 "
            f"통계적으로 유의하였다. 지속성 대 LSTM {pk(sg('overall','Persistence','LSTM')['p_holm'])}"
            f"(r = {sg('overall','Persistence','LSTM')['rank_biserial_r']:.2f}), 지속성 대 1D-CNN "
            f"{pk(sg('overall','Persistence','1D-CNN')['p_holm'])}"
            f"(r = {sg('overall','Persistence','1D-CNN')['rank_biserial_r']:.2f}), XGBoost 대 LSTM "
            f"{pk(sg('overall','XGBoost','LSTM')['p_holm'])}"
            f"(r = {sg('overall','XGBoost','LSTM')['rank_biserial_r']:.2f}), XGBoost 대 1D-CNN "
            f"{pk(sg('overall','XGBoost','1D-CNN')['p_holm'])}"
            f"(r = {sg('overall','XGBoost','1D-CNN')['rank_biserial_r']:.2f})이었다. LSTM–1D-CNN "
            f"차이는 보정 후 유의하지 않았다({pk(sg('overall','LSTM','1D-CNN')['p_holm'])}).")

    heading(d, "3.4. 발생 예측 평균 성능", level=2)
    para(d, "No-Onset 기준모델은 모든 지역에서 F1과 재현율이 0이었다. 양성 예측을 전혀 생성하지 "
            "않으므로 정밀도는 정의되지 않으며 0으로 처리하였다. 이는 발생 후보집합(현재 비안개 "
            "사례) 내에서 발생 사례가 극히 희소하다는 사실을 반영한다. 모든 사례를 비발생으로 "
            "예측하면 높은 정확도를 얻지만 조기경보의 실제 대상인 실제 발생 사례는 단 한 건도 "
            "탐지하지 못한다.")
    para(d, f"XGBoost-Onset이 가장 높은 평균 F1({ms(f1n['XGBoost-Onset'])})과 평균 PR-AUC "
            f"{prn['XGBoost-Onset'].mean():.3f}을 기록하였고, LSTM-Onset"
            f"(F1 {ms(f1n['LSTM-Onset'])}, PR-AUC {prn['LSTM-Onset'].mean():.3f})과 1D-CNN-Onset"
            f"(F1 {ms(f1n['1D-CNN-Onset'])}, PR-AUC {prn['1D-CNN-Onset'].mean():.3f})이 뒤를 "
            f"이었다. XGBoost는 12개 중 {xgb_best_onset}개 관측소에서 최고 F1을 기록하였으며, "
            f"예외는 {', '.join(losers_ko)}로 이들 관측소에서는 딥러닝 모델이 더 높았다. 평균 "
            "PR-AUC에서는 XGBoost와 LSTM의 순서가 사실상 역전되므로, 발생 탐지에서 XGBoost의 우위는 "
            "일관되지만 균일하지는 않다고 기술해야 한다.")
    para(d, "12개 관측소에 대한 Wilcoxon 부호순위 검정도 이를 확인한다. XGBoost는 F1에서 LSTM보다 "
            f"유의하게 우수하지 않았고({pk(sg('onset','XGBoost-Onset','LSTM-Onset')['p_holm'])}), "
            f"PR-AUC에서는 근소하게 뒤졌다"
            f"({pk(sg('onset','XGBoost-Onset','LSTM-Onset','pr_auc')['p_holm'])}). 1D-CNN 대비 F1 "
            f"우위도 Holm 보정을 통과하지 못하였다"
            f"({pk(sg('onset','XGBoost-Onset','1D-CNN-Onset')['p_holm'])}). 다만 PR-AUC에서는 "
            f"XGBoost와 LSTM 모두 1D-CNN을 유의하게 앞섰다"
            f"({pk(sg('onset','XGBoost-Onset','1D-CNN-Onset','pr_auc')['p_holm'])} 및 "
            f"{pk(sg('onset','LSTM-Onset','1D-CNN-Onset','pr_auc')['p_holm'])}). 따라서 방어 "
            "가능한 진술은 정형 모델과 순환 모델이 발생 탐지에서 비슷한 성능을 보이며, 두 모델 모두 "
            "합성곱 모델보다 희소 양성 클래스를 더 잘 순위화한다는 것이다. 모든 학습 모델에서 절대 "
            "성능은 낮게 유지되어 이 과제의 어려움을 보여준다.")
    rows = [["모델", "F1", "PR-AUC", "비고"]]
    rows.append(["No-Onset 기준모델", "0.000", "—", "구조적으로 탐지 불가"])
    for mdl, ko in [("XGBoost-Onset", "XGBoost-Onset"), ("1D-CNN-Onset", "1D-CNN-Onset"),
                    ("LSTM-Onset", "LSTM-Onset")]:
        rows.append([ko, f"{f1n[mdl].mean():.3f} ± {f1n[mdl].std(ddof=1):.3f}",
                     f"{prn[mdl].mean():.3f}",
                     f"12개 중 {int((best_onset==mdl).sum())}개 관측소 최고"])
    caption(d, "표 4. 12개 지역 평균 발생 예측 성능.")
    table(d, rows, widths=[4.4, 3.6, 2.8, 5.4])
    d.add_paragraph()
    figure(d, "image4.png", f"그림 4. 시험셋 발생률 순으로 정렬한 관측소별 0→1 안개 발생 탐지 성능 "
                            f"비교. XGBoost-Onset이 12개 중 {xgb_best_onset}개 관측소에서 최고 F1을 "
                            f"기록하였다.")
    para(d, "지속성 기준모델이 가장 강했던 전체 T+1 예측과 달리, 발생 예측은 지속성이 전혀 다룰 수 "
            "없는 과제이며 여기서는 학습 모델만이 유일한 신호를 제공한다. XGBoost가 가장 높은 평균 "
            f"F1과 12개 중 {xgb_best_onset}개 관측소 최고 F1을 기록한 반면 LSTM은 평균 PR-AUC가 "
            "근소하게 높았으며, 두 쌍별 차이 모두 통계적으로 유의하지 않았다. 단일 관측소 정형 "
            "기상자료에서 극히 희소하고 비선형적인 전이 신호를 학습하는 데 있어, 본 실험 구성에서는 "
            "트리 기반 앙상블과 순환 모델이 비슷한 선택지로 보인다.")

    heading(d, "3.5. TP·FP·FN별 대상 시각 실제 시정 분포", level=2)
    figure(d, "image5.png", "그림 5. TP, FP, FN 사례의 대상 시각 실제 시정 분포. 전체 T+1 XGBoost "
                            "모델 기준이며, 파선은 안개를 정의하는 시정 1 km 임계값이다.", 11.5)
    para(d, f"이 분석은 발생 전용 모델이 아니라 전체 T+1 XGBoost 모델을 사용하여 수행하였다. TP "
            f"사례(n = {n_tp:,})의 대상 시각 시정 중앙값은 {tp:.2f} km인 반면 FN 사례"
            f"(n = {n_fn:,})는 {fn:.2f} km로 더 높았다. TP와 FN 모두 실제 안개 조건에 해당하므로, "
            f"미탐지 사례는 정탐지 사례보다 대상 시각의 시정 저하가 덜 심각한 경향이 있음을 "
            f"의미한다. FP 사례(n = {n_fp:,})의 대상 시각 시정 중앙값은 {fp:.2f} km이며 분포가 "
            "상당히 넓었다. 따라서 모델이 양성으로 판단한 조건이 항상 시정 1 km 미만에 해당하지는 "
            "않았다. 본 분석은 전체 T+1 과제의 모든 전이 유형을 포함하므로 FP 분포를 안개 발생 예측 "
            "실패로만 해석해서는 안 된다.")

    heading(d, "3.6. 지역별 SHAP 변수 중요도 및 관측소 간 일관성", level=2)
    figure(d, "image6.png", "그림 6a. 전체 T+1 XGBoost 모델의 전역 변수 중요도. 12개 ASOS 관측소의 "
                            "관측소별 평균 절대 SHAP 값을 평균하여 산출하였다.", 12.5)
    figure(d, "image7.png", "그림 6b. 관측소별 정규화 SHAP 값(행합 정규화). 색이 진할수록 해당 "
                            "관측소 내 상대적 중요도가 크다.")
    para(d, f"그림 6a는 12개 관측소의 평균 절대 SHAP 값을 평균하여 산출한 XGBoost 모델의 전역 변수 "
            f"중요도를 보여준다. 현재 시정이 압도적으로 가장 높은 기여도를 보였고"
            f"(평균 |SHAP| {shap_g['mean_abs_shap'].iloc[0]:.3f}), 이슬점차"
            f"({shap_g['mean_abs_shap'].iloc[1]:.3f}), 직전 시각 시정"
            f"({shap_g['mean_abs_shap'].iloc[2]:.3f}), 상대습도"
            f"({shap_g['mean_abs_shap'].iloc[3]:.3f})가 뒤를 이었다. 이는 모델이 안개 발생을 예측할 "
            "때 주로 현재 시정 상태와 대기 포화 조건에 의존했음을 나타낸다. 특히 현재 시정의 지배적 "
            "기여는 안개와 시정 악화의 강한 시간적 지속성을 반영하며, 바로 그 변수에 기반한 지속성 "
            "기준모델을 능가하기 어렵다는 결과와도 일관된다.")
    para(d, "관측소별 변수 중요도 구조를 더 살펴보기 위해, 각 관측소에서 변수별 평균 |SHAP| 값을 "
            "해당 관측소의 모든 평균 |SHAP| 값의 합으로 나눈 정규화 SHAP 값을 계산하였다(그림 6b). "
            "그림 6b에서 보듯 변수 중요도 구조는 12개 관측소 전반에서 대체로 유사하였다. 현재 시정, "
            "이슬점차, 습도가 대부분의 관측소에서 가장 높은 상대 기여도를 보였고, 직전 시정, 시각, "
            "기압, 풍속, 지면–대기 온도차 등 보조 변수의 기여는 상대적으로 작았다. 이는 모델이 "
            "관측소별로 뚜렷이 구분되는 변수 중요도 구조를 보이기보다, 관측소 전반에서 대체로 "
            "유사한 예측변수 집합에 의존했음을 의미한다.")
    para(d, "관측소 간 수치적 차이는 일부 존재했으나 지역 유형별로 명확히 구분되는 변수 중요도 "
            "패턴은 확인되지 않았다. 따라서 SHAP 결과는 12개 관측소 전체에서 동일한 안개 형성 기구가 "
            "작동한다는 증거라기보다, 모델 의존 구조의 관측소 간 유사성에 대한 증거로 해석하는 것이 "
            "타당하다. 다만 이 결과는 현재 입력변수 집합이 관측소별 안개 형성 기구의 차이를 충분히 "
            "분리하지 못한다는 한계도 함께 드러낸다. 지형, 해륙 대비, 복사, 운량, 종관 규모 흐름 "
            "같은 지역 특이 변수를 포함하지 않았기 때문에 지역별 안개 유형의 차이가 SHAP 값에 "
            "명확히 반영되지 않았을 수 있다.")

    # ---------------- 4. 논의 ----------------
    heading(d, "4. 논의", level=1)
    heading(d, "4.1. 전체 T+1 예측에서 지속성과 XGBoost의 경쟁적 성능", level=2)
    para(d, f"전체 T+1 예측에서 지속성은 XGBoost보다 높은 평균 F1을 기록하였고"
            f"({f1o['Persistence'].mean():.3f} 대 {f1o['XGBoost'].mean():.3f}, 차이 "
            f"{f1o['Persistence'].mean()-f1o['XGBoost'].mean():.3f}), 검증셋 기반 임계값 조정 "
            f"후에도 12개 중 {pers_wins}개 관측소에서 더 우수하였다. 이는 학습 모델이 시정 지속성 "
            "이상의 정보를 추출하고 있다면 기대할 수 있는 순서와 반대이다. 지속성의 우위는 발생률이 "
            "낮은 관측소에 집중되어 있는데, 이는 해당 관측소에서 이용 가능한 양성 학습 표본이 "
            "제한적이며 그것이 학습된 결정경계의 견고성에 미친 영향을 부분적으로 반영할 수 있다. "
            "발생률이 가장 높은 세 관측소에서는 두 모델이 근접하였다.")
    para(d, f"12개 관측소에 대한 Wilcoxon 부호순위 검정 결과는 보정 전 "
            f"{pk(sg('overall','Persistence','XGBoost')['p_raw'])}, Holm 보정 후 "
            f"{pk(sg('overall','Persistence','XGBoost')['p_holm'])}로, 이 차이는 확립된 것이 아니라 "
            "시사적인 수준이다. 따라서 본 연구가 도출하는 결론은 의도적으로 약한 쪽이다. XGBoost는 "
            "1시간 후 안개 상태 예측에서 시정 지속성 대비 개선을 입증하지 못하였으며, 안개가 희소한 "
            "지역에서는 뚜렷하게 더 나빴다. 예보 실무 관점에서 함의는 어느 쪽이든 동일하다. 지속성은 "
            "학습도, 임계값 조정도, 조건 변화에 따른 재학습도 필요하지 않기 때문이다.")
    para(d, f"한편 분류 임계값을 기본값 0.5로 고정하면 지속성이 12개 전 관측소에서 XGBoost보다 "
            f"우수하며 XGBoost 평균 F1은 {f1o['XGBoost'].mean():.3f}에서 {xgb_def_mean:.3f}로 "
            "하락한다. 이는 학습 모델의 확률 출력이 기본 임계값에서 최적 성능을 보장하지 않으며, "
            "안개 예측처럼 클래스 불균형이 심한 문제에서는 임계값 선택 절차가 모델 성능 평가에 "
            "상당한 영향을 줄 수 있음을 시사한다. Shin 등[19]도 클래스 불균형과 시간적 분포 변화가 "
            "국내 시정 초단기예측에서 성능 저하의 주요 요인이라고 지적하였다. 따라서 단기 안개 예측 "
            "모델 비교에서는 알고리즘 구조뿐 아니라 검증 기반 임계값 최적화 적용 여부와 선택된 "
            "평가지표를 명확히 보고해야 한다.")
    para(d, "아울러 전체 T+1 예측에서 지속성의 강한 성능이 곧 강한 조기경보 성능을 의미하지는 "
            "않는다. 지속성은 현재 상태가 다음 시각에도 이어진다고 가정하므로, 비안개에서 안개로의 "
            "전이인 0→1 발생을 구조적으로 탐지할 수 없다. Schütz 등[1]도 전체 시정 예측 성능만으로는 "
            "안개 형성·소멸 같은 전이 사건에 대한 예측 성능이 가려질 수 있다고 지적하였다. 따라서 "
            "전체 T+1 예측과 발생 예측은 분리하여 해석해야 한다.")

    heading(d, "4.2. 발생 탐지: 정형 모델과 순환 모델의 비슷한 성능", level=2)
    para(d, f"발생 예측에서 XGBoost는 12개 중 {xgb_best_onset}개 관측소에서 1D-CNN과 LSTM보다 높은 "
            f"F1을 보였고 가장 높은 평균 F1({f1n['XGBoost-Onset'].mean():.3f})을 기록하였다. "
            f"LSTM은 {f1n['LSTM-Onset'].mean():.3f}, 1D-CNN은 "
            f"{f1n['1D-CNN-Onset'].mean():.3f}이었다. 그 차이는 전체 과제에서 지속성과 학습 모델을 "
            f"가르는 격차보다 작고 균일하지 않으며, 평균 PR-AUC에서는 LSTM이 근소하게 앞섰다"
            f"({prn['LSTM-Onset'].mean():.3f} 대 {prn['XGBoost-Onset'].mean():.3f}). 따라서 본 "
            "자료 구성에서 정형 모델과 순환 모델은 0→1 발생 탐지에서 비슷하며, XGBoost가 F1에서 "
            "작고 비유의한 우위를, LSTM이 PR-AUC에서 작고 비유의한 우위를 갖는다고 진술할 수 있다. "
            "XGBoost가 딥러닝 모델을 압도한다는 기존 주장도, 반대로 LSTM에 유리한 주장도 12개 대응 "
            "관측치로는 뒷받침되지 않는다.")
    para(d, f"첫째, 발생 사례는 전체 시험 사례의 약 {onset_pct:.2f}%에 불과한 극소수 클래스이다. "
            "이러한 조건에서 모델은 다수 비안개 클래스로 편향되어 높은 정확도를 얻으면서도 실제 안개 "
            "발생을 탐지하지 못하기 쉽다. Shin 등[19]과 Kim 등[20] 역시 희소성과 분포 변화가 성능 "
            "저하의 핵심 요인이라고 설명하였다.")
    para(d, "둘째, XGBoost는 scale_pos_weight 조정과 결합하여 개별 변수에 대한 비선형 분할 경계를 "
            "직접 학습할 수 있어 희소 신호에 상대적으로 민감할 수 있다. 이러한 선행 연구 결과들은 "
            "정형 기상자료를 이용한 안개 예측에서 트리 기반 앙상블의 지속적인 실용성과 경쟁력을 "
            "뒷받침하지만, 이들이 순차 모델을 능가함을 의미하지는 않는다.")
    para(d, "셋째, 본 연구에서 사용한 물리 파생변수(이슬점차, 직전 시정, 시차 변수)는 이미 단기 "
            "변화 정보를 압축적으로 부호화하고 있다. XGBoost는 이러한 변수를 분할 기준으로 직접 사용할 "
            "수 있어 특정 시점의 급격한 조건 변화가 중요한 발생 예측 과제에서 효과적일 수 있다. "
            "반면 1D-CNN과 LSTM은 6시간 시계열 윈도우 내에서 패턴을 학습해야 하므로 발생처럼 희소하고 "
            "지속 시간이 짧은 전이 신호가 순차 구조 안에서 희석될 수 있다. 이는 딥러닝 모델이 안개 "
            "예측에 본질적으로 부적합함을 의미하지 않는다. NWP 산출물, 위성 기반 해수면온도, 3차원 "
            "격자 입력을 활용하는 Kamangir 등[22]의 FogNet 같은 모델은 현업 앙상블을 능가할 수 있다. "
            "개별 ASOS 관측소의 정형 자료와 짧은 시계열 윈도우라는 본 연구의 설정에서는 XGBoost와 "
            "LSTM이 비슷한 성능을 보였고, 관측소 수준 검정으로는 둘 사이의 차이를 판별하지 못하였다.")

    heading(d, "4.3. SHAP 기반 해석 가능성과 지역 일관성", level=2)
    para(d, "SHAP 분석 결과 현재 시정, 이슬점차, 상대습도가 12개 관측소 전반에서 상위에 위치하였다. "
            "이는 XGBoost 모델이 관측소 전반에서 대체로 유사한 예측변수 위계에 의존했음을 나타낸다. "
            "다만 변수 중요도의 유사성이 곧 지역별 안개 형성 기구가 동일함을 의미하지는 않는다. "
            "시정·습도 관련 변수의 중요성은 물리적으로도 타당하다. Gultepe와 Milbrandt[10]는 상대습도, "
            "강수, 시정 사이의 관계에 기반한 확률적 시정 모수화를 제안하여 시정 악화가 수분 조건과 "
            "밀접히 연관됨을 보였다. Long 등[17]도 상대습도, 액체수함량, 물방울 수농도, 입경, "
            "소산계수를 안개 시정 모수화의 핵심 인자로 제시하였다. 본 연구는 ASOS 관측의 한계로 "
            "미세물리 변수를 직접 포함하지 못하였으나, 이슬점차·상대습도·직전 시정은 이러한 물리 "
            "과정과 연관된 대기 조건의 간접 지표로 기능할 수 있다.")
    para(d, "선행 연구들은 한반도 지역별로 안개 형성 기구가 뚜렷이 다르다고 보고한다. 본 연구에서 "
            "SHAP 값의 뚜렷한 지역 차이가 나타나지 않은 것은 지역 특성이 없다는 의미가 아니라, "
            "입력자료가 대체로 개별 관측소의 지상 기상변수에 국한되어 지역별 안개 기구를 충분히 "
            "분리하지 못했음을 반영할 가능성이 크다. 향후 연구에서는 고도, 지형 기복, 해안 거리, "
            "토지피복, 복사, 운량, 바람 특성, 대기 안정도, 해수면온도, 인접 관측소 정보 등 지역 "
            "특성을 설명할 수 있는 변수를 추가해야 한다.")

    heading(d, "4.4. 한계", level=2)
    para(d, "본 연구의 한계는 다음과 같다. 첫째, 두 모델군에 동일한 시간 분할, 동일한 평가 사례, "
            "동일한 정보 지평을 부여하였으나 입력 표현과 초매개변수 탐색 범위는 완전히 동일하지 "
            "않으므로, 본 비교는 순수한 구조 비교가 아니라 실용적 모델 구성 간의 비교로 해석해야 "
            "한다. 본 연구의 딥러닝 모델은 6시간 ASOS 기반 시퀀스 입력을 사용하는 1D-CNN과 LSTM에 "
            "국한되며, NWP 산출물·위성·해수면온도·3차원 격자 입력을 활용하는 FogNet 같은 고차원 "
            "구조와는 문제 설정이 다르다[22]. 따라서 본 결과는 딥러닝이 안개 예측에 부적합하다는 "
            "증거로 읽혀서는 안 되며, 본 자료 구성에서 두 딥러닝 모델 모두 학습이 필요 없는 지속성 "
            "기준모델을 개선하지 못하였다는 의미로 이해해야 한다.")
    para(d, "둘째, 통계적 검정력이 제한적이다. 12개의 대응 관측소 관측치로는 양측 Wilcoxon 부호순위 "
            "검정이 크고 일관된 차이를 탐지하는 데 적합하며, 실제로 지속성–XGBoost 및 XGBoost–LSTM "
            "비교는 점추정치가 다름에도 Holm 보정 후 유의성에 도달하지 못하였다. 이러한 비유의 "
            "결과는 현재 자료로 그 차이를 판별할 수 없다는 뜻이며, 모델들이 동등함을 입증하는 것은 "
            "아니다. 또한 12개 관측소가 동일한 지역 기상계 내에 지리적으로 분포하므로 완전히 독립적인 "
            "공간 단위를 구성하지 않을 수 있다는 점에서, 관측소 수준 검정은 신중히 해석해야 한다. "
            "보고된 관측소 간 표준편차(±)는 추정의 불확실성이 아니라 관측소 간 이질성을 반영한다. "
            "따라서 모델 간의 작은 F1 차이는 확정적 순위가 아니라 경향으로 해석해야 하며, 본 연구는 "
            "검정이 뒷받침하지 않는 순위 주장을 의도적으로 배제하였다. 끝으로 1D-CNN과 LSTM 결과는 "
            "무작위 초기화에 의존하며, 보고된 값은 고정 시드 단일 실행 결과이고 시드 간 변동성은 "
            "공개 저장소에 보고하였다.")
    para(d, "셋째, 인제·서울·속초처럼 안개 발생률이 극히 낮은 지역에서는 양성 표본 수가 적어 모든 "
            "모델에서 성능 추정과 SHAP 해석의 신뢰도가 제약되었다. 넷째, 임계값은 단일 검증연도"
            "(2022년)로 최적화하였다. 시험셋 평가 전에 고정하기는 하였으나 연간 안개 빈도와 기상 "
            "조건 변화에서도 최적으로 유지된다고 보장할 수 없다. 다섯째, 개별 관측소 수준 입력변수만 "
            "사용하여 인접 관측소 간 공간 정보, 해양 조건, 지형 구조, 위성 기반 운/해무 정보, NWP "
            "기반 경계층 정보를 충분히 반영하지 못하였다. 여섯째, SHAP 분석은 주요 입력변수 해석에 "
            "유용하나 안개 미세물리 과정을 직접 설명하지는 못한다. 마지막으로 본 연구는 전체 T+1 "
            "예측과 발생 예측을 분리하였으나 안개 지속시간이나 소멸 시점을 직접 예측하지는 않았다.")

    # ---------------- 5. 결론 ----------------
    heading(d, "5. 결론", level=1)
    para(d, f"전체 T+1 예측에서는 학습이 필요 없는 지속성 기준모델이 가장 높은 평균 F1"
            f"({ms(f1o['Persistence'])})을 기록하였고 XGBoost({ms(f1o['XGBoost'])}), "
            f"LSTM({ms(f1o['LSTM'])}), 1D-CNN({ms(f1o['1D-CNN'])})이 뒤를 이었다. 지속성은 임계값 "
            f"조정 후 12개 중 {pers_wins}개 관측소에서, 기본 임계값에서는 12개 전 관측소에서 "
            f"XGBoost보다 우수하였으며 안개가 희소해질수록 격차가 벌어졌다. 다만 XGBoost 대비 "
            f"우위는 12개 관측소에서 통계적으로 유의하지 않았으므로"
            f"(Holm 보정 후 {pk(sg('overall','Persistence','XGBoost')['p_holm'])}), 본 연구는 "
            "XGBoost가 기준모델 대비 개선을 입증하지 못하였다는 결론만을 도출한다. 두 딥러닝 모델 "
            "대비 우위는 본 연구에서 평가한 전체 T+1 과제에서 유의하였다. 결과적으로 1시간 후 안개 "
            "상태 예측에서 학습이 필요 없는 지속성 기준모델은 본 연구의 어떤 학습 모델에도 능가당하지 "
            "않았으며, 동시에 0→1 안개 발생을 구조적으로 탐지할 수 없다.")
    para(d, f"지속성이 구조적으로 적용 불가능한 발생 탐지에서는 XGBoost가 가장 높은 평균 F1"
            f"({ms(f1n['XGBoost-Onset'])})을 기록하고 12개 중 {xgb_best_onset}개 관측소에서 "
            f"최고였으며, LSTM({ms(f1n['LSTM-Onset'])})과 1D-CNN({ms(f1n['1D-CNN-Onset'])})이 뒤를 "
            "이었다. 평균 PR-AUC에서는 LSTM이 근소하게 앞섰고, XGBoost와 LSTM 사이의 F1 및 PR-AUC "
            "차이는 모두 통계적으로 유의하지 않았다. 모든 학습 모델의 절대 성능이 낮게 유지되어 안개 "
            "발생이 전체 안개 상태 예측보다 훨씬 어려운 과제임을 보여준다. 본 ASOS 기반 실험 구성에서 "
            "XGBoost와 LSTM은 본 연구의 통계적 해상도 내에서 비슷한 발생 탐지 성능을 보였으며, 어느 "
            "구조에도 일반적 우위를 뒷받침하는 통계적 근거는 없다.")
    para(d, "SHAP 분석 결과 현재 시정, 이슬점차, 직전 시각 시정, 상대습도가 일관되게 가장 영향력 있는 "
            "예측변수로 나타났다. 현재 시정의 지배적 기여는 지속성 기준모델의 강력함과 일관된다. "
            "관측소 간 변수 순위의 유사성은 모델 의존 구조의 공통 패턴을 나타내지만 지역별 안개 형성 "
            "기구가 동일함을 입증하지는 않는다. 이러한 결과를 종합하면, 전체 1시간 후 안개 상태 예측 "
            "성능이 우수하다고 해서 그것이 곧 유용한 안개 발생 조기경보 능력으로 이어지지는 않는다. "
            "전체 과제에서 가장 이기기 어려웠던 기준모델이 바로 발생 경보를 전혀 낼 수 없는 모델이기 "
            "때문이다. 따라서 전체 안개 상태 예측과 안개 발생 탐지는 서로 다른 과제로 규정하고 별도로 "
            "평가·보고해야 하며, 본 연구의 기여는 새로운 최고성능 모델이 아니라 평가 프레임워크와 "
            "진단적 발견이다. 향후 연구는 공간·지형·복사·해양·경계층 정보를 반영하고, 시간적 분포 "
            "변화 아래에서 임계값 보정과 희소사건 학습을 평가하며, 안개 지속과 소멸 예측으로 이 "
            "프레임워크를 확장해야 한다.")

    # ---------------- back matter ----------------
    heading(d, "자료 이용 가능성", level=1)
    para(d, "본 연구에 사용된 2015–2024년 시간별 ASOS 관측자료는 기상자료개방포털"
            "(https://data.kma.go.kr)에서 공개적으로 이용할 수 있다. 본 논문의 분석, 표, 그림을 "
            "재현하는 데 필요한 모든 소스코드, 설정파일, 결과파일은 원시 관측자료의 내려받기 및 준비 "
            "절차와 함께 공개 저장소에 공개되어 있다. 저장소는 원시 관측자료로부터 보고된 모든 지표, "
            "통계검정 결과, 표, 그림을 재생성하는 문서화된 파이프라인을 제공한다. 학습된 모델 파일은 "
            "파이프라인이 직접 재현하므로 별도로 배포하지 않는다.")
    heading(d, "참고문헌", level=1)
    para(d, "참고문헌 목록은 영문 정본(Atmosphere_eng_v8_final.docx)의 References 절을 참조.",
         size=9.5, italic=True)

    d.save(out_path)
    return out_path


# ================================================================ COVER LETTER
def build_cover_letter(out_path):
    d = new_doc()
    para(d, "[검토용 국문본]", size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    heading(d, "Cover Letter (국문 검토본)", level=0)
    notice(d, [
        "■ 본 문서는 지도교수 검토를 위한 국문 번역본입니다. 실제 제출본은 "
        "cover_letter_v4_final.docx(영문)이며, 표현이 다를 경우 영문본이 우선합니다.",
        "■ 모든 수치는 영문본과 동일하게 결과 CSV에서 자동으로 읽어 생성되었습니다.",
    ])

    para(d, "Atmosphere 편집위원장께,")
    para(d, "저희 논문 「Short-term Fog Forecasting and Onset Detection Using Multi-station ASOS "
            "Observations in Korea: A Comparison of Persistence, XGBoost, and Deep Learning "
            "Models」를 Atmosphere에 게재 신청하고자 제출합니다.")

    para(d, "연구 동기 및 기여.", bold=True, space_after=2)
    para(d, "안개 및 저시정 예측에 관한 최근 기계학습 문헌의 상당수는 단순한 지속성 기준모델과의 "
            "비교 없이, 그리고 일상적인 안개 지속과 운영상 결정적인 안개 발생을 분리하지 않은 채 "
            "높은 종합 성능 지표를 보고합니다. 본 연구는 최고 성능 달성을 주장하기보다 방법론적·진단적 "
            "기여로 의도적으로 구성되었습니다. 한반도 중북부 12개 관측소의 10년(2015–2024) 시간별 "
            "ASOS 관측자료를 이용하여, 연구자와 예보관 모두에게 유용하다고 판단되는 세 가지 논점을 "
            "제시합니다.")

    para(d, "1. 전체 T+1 안개 예측과 0→1 발생 예측을 분리하여, 전체 1시간 후 안개 상태 예측 성능이 "
            "우수하다고 해서 그것이 곧 유용한 안개 발생 조기경보 능력으로 이어지지는 않음을 보였습니다. "
            "전체 과제에서 가장 이기기 어려운 기준모델이 바로 발생 경보를 전혀 낼 수 없는 모델입니다. "
            "따라서 저희의 핵심 주장은 두 과제를 서로 다른 과제로 규정하고 별도로 평가해야 한다는 "
            "것입니다.")

    para(d, f"2. 지속성이 초단기 안개 예측에서 매우 강력한 기준모델임을 보였습니다. 12개 관측소에서 "
            f"비교한 네 모델 중 가장 높은 평균 F1을 기록하였고({ms(f1o['Persistence'])} 대 XGBoost "
            f"{ms(f1o['XGBoost'])}), 검증셋 기반 임계값 조정 후 12개 중 {pers_wins}개 관측소에서, "
            f"기본 결정 임계값에서는 12개 전 관측소에서 XGBoost를 앞섰습니다. 안개가 희소한 곳일수록 "
            "격차가 커졌는데 이는 해당 관측소의 양성 학습 사례 희소성과 부합하는 양상입니다. 다만 12개 "
            f"대응 관측치에서 지속성–XGBoost 차이는 Holm 보정 후 유의하지 않으므로"
            f"({pk(sg('overall','Persistence','XGBoost')['p_holm'])}), 저희는 XGBoost가 지속성을 "
            f"개선하지 못하였다고만 주장하며 지속성이 우월하다고 주장하지 않습니다. 다만 두 모델 모두 "
            f"본 실험 구성에서 두 딥러닝 모델은 유의하게 앞섰습니다"
            f"({pk(sg('overall','XGBoost','LSTM')['p_holm'])} 이하). 그럼에도 지속성은 구조적으로 "
            "발생을 탐지할 수 없습니다. 이는 학습 기반 모델이 무엇을 더하고 무엇을 더하지 못하는지를 "
            "분명히 해 주며, 이런 종류의 명확한 부정적 결과도 보고할 가치가 있다고 생각합니다.")

    para(d, f"3. 발생 탐지가 시험한 모든 모델에서 절대적으로 여전히 어렵고 대체로 해결되지 않은 "
            f"과제임을 보였습니다. 트리 기반 모델(XGBoost, 평균 F1 {f1n['XGBoost-Onset'].mean():.3f})이 "
            f"가장 높은 평균 F1을 기록하고 12개 중 {xgb_best_onset}개 관측소에서 최고였으며 "
            f"LSTM({f1n['LSTM-Onset'].mean():.3f})과 1D-CNN({f1n['1D-CNN-Onset'].mean():.3f})이 뒤를 "
            f"이었습니다. 평균 PR-AUC에서는 LSTM이 근소하게 앞섰고"
            f"({prn['LSTM-Onset'].mean():.3f} 대 {prn['XGBoost-Onset'].mean():.3f}) 두 차이 모두 "
            f"통계적으로 유의하지 않으므로"
            f"(F1 {pk(sg('onset','XGBoost-Onset','LSTM-Onset')['p_holm'])}, "
            f"PR-AUC {pk(sg('onset','XGBoost-Onset','LSTM-Onset','pr_auc')['p_holm'])}), 저희는 두 "
            "모델을 본 자료의 통계적 해상도 내에서 비슷하다고 보고하며 순위를 매기지 않습니다. 두 "
            "구조가 동등하다고 주장하지도 않습니다. SHAP 분석은 예측이 물리적으로 해석 가능하고 지역 "
            "간 일관된 변수(현재 시정, 이슬점차, 직전 시정, 상대습도)에 의해 주도됨을 보여줍니다.")

    para(d, "한계의 명시.", bold=True, space_after=2)
    para(d, "본문에 한계를 명확히 기술하였습니다. XGBoost와 딥러닝 모델의 비교는 통제된 구조 비교가 "
            "아니라 실용적 모델 구성 간 비교이나, 시간 분할·평가 사례·모든 모델이 이용 가능한 정보 "
            "지평은 동일하게 맞추었습니다. 모델 차이는 12개 관측소에 대한 양측 Wilcoxon 부호순위 "
            "검정과 Holm 보정으로 검정하였으며, 검정이 뒷받침하지 않는 순위 주장은 의도적으로 "
            "배제하였습니다. 12개 대응 관측소만으로는 검정력이 제한적이고, 비유의 결과가 동등성을 "
            "입증하지 않으며, 관측소들이 동일한 지역 기상계 내에 있어 완전히 독립적인 공간 단위를 "
            "구성하지 않을 수 있습니다. 보고된 관측소 간 표준편차는 추정 불확실성이나 실행 간 안정성이 "
            "아니라 지역 이질성을 반영합니다. 발생률이 매우 낮은 관측소는 불확실성이 크며, 딥러닝 "
            "결과는 무작위 초기화에 의존하므로 고정 시드 실행 결과를 보고하고 시드 간 변동성은 "
            "저장소에 정량화하였습니다. 이러한 부정적이고 경고적인 결과를 정직하게 보고하는 것 자체가 "
            "기여라고 생각하며, 이는 현업 안개 초단기예측에서 종합 지표의 과대 해석을 방지하는 데 "
            "도움이 됩니다.")

    para(d, "본 원고는 독창적이며 다른 곳에 게재된 바 없고 다른 학술지에서 심사 중이지 않습니다. 모든 "
            "저자가 투고에 동의하였으며 이해상충은 없습니다. 사용된 ASOS 관측자료는 기상청에서 공개 "
            "제공됩니다. 검증을 지원하기 위해 전체 파이프라인, 설정파일, 결과파일을 담은 공개 저장소를 "
            "공개합니다. 문서화된 파이프라인이 원시 관측자료로부터 보고된 모든 지표, 통계검정 결과, "
            "표, 그림을 재생성하며, 딥러닝 모델은 매 실행 전에 1D-CNN과 LSTM이 기술된 구조와 실제로 "
            "일치하는지 구조적으로 점검됩니다. 학습된 모델 파일은 파이프라인이 원시 자료로부터 "
            "재현하므로 별도 배포하지 않습니다.")
    para(d, "심사를 맡아 주셔서 감사드리며 심사위원들의 의견을 기다리겠습니다.")
    para(d, "감사합니다.", space_after=14)
    para(d, "[교신저자명], 저자 일동을 대표하여")
    para(d, "[소속]")
    para(d, "[이메일]")

    d.save(out_path)
    return out_path


if __name__ == "__main__":
    a = build_manuscript(os.path.join(HERE, "안개예측_논문_국문검토본_v8.docx"))
    b = build_cover_letter(os.path.join(HERE, "안개예측_커버레터_국문검토본_v4.docx"))
    print("Wrote:", a)
    print("Wrote:", b)
